"""Generate 红队专项攻击支撑方案 document using default proposal template."""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

TEMPLATE = Path("templates/default-proposal-template.docx")
OUTPUT = Path("output/红队专项攻击支撑方案.docx")

BODY_FIRST_LINE_INDENT = Pt(28)
BODY_FONT_SIZE = Pt(11)
TABLE_FONT_SIZE = Pt(9)

# ── helpers ──────────────────────────────────────────────────────────


def set_run_font(run, size=BODY_FONT_SIZE, bold=None, color=None):
    run.font.size = size
    if bold is not None:
        run.bold = bold
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    if color:
        run.font.color.rgb = color


def set_paragraph_font(paragraph, size=BODY_FONT_SIZE, bold=None):
    for run in paragraph.runs:
        set_run_font(run, size=size, bold=bold)


def style_body(paragraph, first_line=True):
    paragraph.paragraph_format.first_line_indent = (
        BODY_FIRST_LINE_INDENT if first_line else None
    )
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.28
    set_paragraph_font(paragraph)


def add_heading_text(doc: Document, text: str, level: int):
    """Add a heading with manual numbering in the text."""
    p = doc.add_heading(text, level=level)
    p.paragraph_format.space_before = Pt(12 if level <= 3 else 8)
    p.paragraph_format.space_after = Pt(6)
    for run in p.runs:
        run.font.name = "宋体"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    return p


def add_body(doc: Document, text: str):
    p = doc.add_paragraph(text)
    style_body(p)


def add_bullet(doc: Document, text: str):
    p = doc.add_paragraph(f"• {text}")
    p.paragraph_format.left_indent = Pt(18)
    p.paragraph_format.first_line_indent = None
    p.paragraph_format.space_after = Pt(3)
    set_paragraph_font(p)


def add_sub_bullet(doc: Document, text: str):
    p = doc.add_paragraph(f"  - {text}")
    p.paragraph_format.left_indent = Pt(36)
    p.paragraph_format.first_line_indent = None
    p.paragraph_format.space_after = Pt(2)
    set_paragraph_font(p, size=Pt(10.5))


def add_bold_lead(doc: Document, lead: str, rest: str):
    """Bold lead-in sentence followed by normal text."""
    p = doc.add_paragraph()
    style_body(p)
    r1 = p.add_run(lead)
    set_run_font(r1, bold=True)
    r2 = p.add_run(rest)
    set_run_font(r2)


def remove_heading_numbering(doc: Document):
    """Remove auto-numbering from Heading 1-6 styles so manual numbering in text works."""
    for level in range(1, 7):
        try:
            style = doc.styles[f"Heading {level}"]
        except KeyError:
            continue
        pPr = style.element.find(qn("w:pPr"))
        if pPr is None:
            continue
        numPr = pPr.find(qn("w:numPr"))
        if numPr is not None:
            pPr.remove(numPr)


def add_table_2col(doc: Document, headers: list[str], rows: list[list[str]]):
    """Add a 2-column formatted table."""
    table = doc.add_table(rows=1 + len(rows), cols=2)
    table.style = "Table Grid"
    table.autofit = True
    # header
    for ci, h in enumerate(headers):
        cell = table.rows[0].cells[ci]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        set_run_font(r, size=TABLE_FONT_SIZE, bold=True)
        _shade_cell(cell, "D9EAF7")
    # data
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = ""
            p = cell.paragraphs[0]
            r = p.add_run(val)
            set_run_font(r, size=TABLE_FONT_SIZE)
    doc.add_paragraph("")  # spacing


def _shade_cell(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def clear_document(document: Document) -> None:
    body = document._element.body
    for child in list(body):
        if child.tag.endswith("sectPr"):
            continue
        body.remove(child)


# ── content ──────────────────────────────────────────────────────────

h3_counter = 0
h4_counters: dict[int, int] = {}
h5_counters: dict[int, int] = {}


def h3(text: str) -> str:
    global h3_counter
    h3_counter += 1
    h4_counters[h3_counter] = 0
    return f"{h3_counter}.1.1 {text}"


def h4(text: str) -> str:
    global h4_counters
    h4_counters[h3_counter] += 1
    ci = h4_counters[h3_counter]
    return f"{h3_counter}.1.1.{ci} {text}"


def h5(text: str) -> str:
    global h5_counters
    h5_counters.setdefault(h3_counter, 0)
    h5_counters[h3_counter] += 1
    ci = h5_counters[h3_counter]
    return f"{h3_counter}.1.1.{h4_counters[h3_counter]}.{ci} {text}"


def build(doc: Document):
    clear_document(doc)
    remove_heading_numbering(doc)

    for section in doc.sections:
        section.top_margin = Inches(0.9)
        section.bottom_margin = Inches(0.9)
        section.left_margin = Inches(0.9)
        section.right_margin = Inches(0.9)

    # ── Title ───────────────────────────────────────────────────────
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(60)
    title_p.paragraph_format.space_after = Pt(8)
    r = title_p.add_run("红队专项攻击支撑方案")
    r.font.size = Pt(22)
    r.bold = True
    r.font.name = "宋体"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_p.paragraph_format.space_after = Pt(24)
    r = sub_p.add_run("——中国联合网络通信有限公司无锡市分公司")
    r.font.size = Pt(14)
    r.font.name = "宋体"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    add_body(
        doc,
        "本方案针对中国联合网络通信有限公司无锡市分公司（以下简称「采购人」）在市级网络攻防实战中的红队攻击能力提升需求，"
        "围绕市级「网安行动」、互联网众测及渗透技能竞赛等实战场景，提供涵盖攻击战术研判、漏洞验证协同、"
        "攻击路径推演及复盘分析等全流程的红队专项攻击支撑服务方案。"
    )

    # ══════════════════════════════════════════════════════════════════
    # CHAPTER 1
    # ══════════════════════════════════════════════════════════════════
    add_heading_text(doc, h3("项目概述"), 3)

    add_heading_text(doc, h4("项目背景"), 4)
    add_body(
        doc,
        "随着市级网络攻防实战对抗的日益激烈，中国联通无锡市分公司在市级「网安行动」、互联网众测及渗透技能竞赛等"
        "活动中，面临真实高水平网络攻击的严峻挑战。为全面提升公司在网络攻防实战中的红队攻击能力，提升实战得分水平"
        "与漏洞发现能力，亟需引入专业的红队专项攻击支撑服务。"
    )
    add_body(
        doc,
        "当前，网络攻击手段日趋复杂多样，攻击者利用先进的攻击工具和技术，能够快速发现并利用系统漏洞。"
        "传统的安全防护手段已难以应对高水平的定向攻击。红队攻击作为检验和提升安全防御能力的有效手段，"
        "通过模拟真实攻击者的战术、技术和流程，帮助组织发现安全防御体系的薄弱环节。"
    )
    add_body(
        doc,
        "本项目通过询比采购方式引入网络安全渗透能力专项支撑服务，由具备丰富红队攻击实战经验的技术专家团队，"
        "为采购人攻击队提供攻击战术研判、漏洞验证协同、攻击路径推演及复盘分析等专业技术支持服务，"
        "全面提升无锡公司面对真实高水平网络攻击的防御能力。"
    )

    add_heading_text(doc, h4("服务目标"), 4)
    add_bold_lead(
        doc, "提升实战攻击能力：",
        "通过专业红队技术支撑，提升采购人攻击队在市级网络安全实战演练中的攻击效率和漏洞发现能力，"
        "确保在授权范围内完成深度渗透测试与漏洞挖掘验证工作。"
    )
    add_bold_lead(
        doc, "建立战术研判机制：",
        "构建7×24小时远程战术研判支撑体系，为攻击队提供持续的战术分析和路径优化建议，"
        "确保攻击活动的高效推进和及时调整。"
    )
    add_bold_lead(
        doc, "强化技术协同能力：",
        "建立多团队协同漏洞验证机制，实现漏洞从发现、验证到复盘的完整闭环管理，"
        "提升漏洞报告的技术深度与有效性。"
    )
    add_bold_lead(
        doc, "形成复盘改进体系：",
        "通过系统的攻击技术复盘和攻击链还原分析，持续改进攻击战术和方法，形成可复用的攻击知识库和最佳实践。"
    )

    add_heading_text(doc, h4("服务范围"), 4)
    add_body(doc, "本方案的服务范围涵盖以下三类主要场景：")
    add_bold_lead(doc, "市级网安行动红队攻击支撑：",
                  "在无锡市「网安行动」期间，为采购人攻击队提供持续性技术研判与战术推演支撑，"
                  "协助攻击队在授权目标范围内高效完成漏洞发现与利用验证工作。")
    add_bold_lead(doc, "市级网信办众测红队支撑：",
                  "在市级互联网众测活动中，为采购人攻击队提供漏洞挖掘方向指引与漏洞利用可行性验证支持，"
                  "提升漏洞报告的技术深度与有效性。")
    add_bold_lead(doc, "市级工信及小规模渗透竞赛支撑：",
                  "在市级专项渗透技能竞赛中，为采购人参赛队伍提供赛前实战化技术集训与赛中远程战术研判支撑。")

    add_heading_text(doc, h4("服务原则"), 4)
    add_body(doc, "本方案在实施过程中严格遵循以下原则：")
    add_bold_lead(doc, "合法合规原则：",
                  "所有红队攻击行为严格遵守《网络安全法》及相关法律法规，仅限于授权目标范围内，不得触及非授权系统。")
    add_bold_lead(doc, "授权范围内原则：",
                  "所有渗透测试和攻击活动均在采购人书面授权范围内进行，明确测试目标和边界，不越界操作。")
    add_bold_lead(doc, "最小影响原则：",
                  "攻击测试活动以不影响业务系统正常运行和不出数据泄露为前提，严格控制测试手段和范围。")
    add_bold_lead(doc, "保密性原则：",
                  "所有参与项目人员签署保密协议，测试过程中获取的敏感数据和漏洞信息严格保密，不得外泄。")

    # ══════════════════════════════════════════════════════════════════
    # CHAPTER 2
    # ══════════════════════════════════════════════════════════════════
    add_heading_text(doc, h3("攻击战术研判"), 3)

    add_body(
        doc,
        "攻击战术研判是红队攻击支撑的核心环节，通过对目标系统的全面信息收集和分析，结合先进的安全攻防框架，"
        "制定精准有效的攻击战术，并在攻击过程中实时调整和优化，确保攻击活动的高效推进。"
    )

    add_heading_text(doc, h4("攻击面分析与目标情报收集"), 4)
    add_body(
        doc,
        "攻击面分析是攻击战术研判的基础，通过多渠道、多维度、多层次的信息收集手段，构建目标系统的完整攻击面画像，"
        "为后续攻击战术制定和路径推演提供数据支撑。"
    )

    add_heading_text(doc, h5("互联网资产探测与暴露面识别"), 5)
    add_body(
        doc,
        "利用先进网络探测技术和情报收集手段，对采购人授权目标范围内的互联网资产进行全面探测和识别："
    )
    add_bullet(doc, "子域名收集：通过DNS枚举、证书透明度日志、搜索引擎语法（Google Dork）等方式进行子域名收集和发现")
    add_bullet(doc, "IP地址段确认：对目标IP地址段进行存活检测和端口扫描，识别开放端口和运行服务")
    add_bullet(doc, "Web应用识别：对目标Web应用进行指纹识别，确定中间件类型、版本及CMS框架信息")
    add_bullet(doc, "SSL/TLS证书分析：通过证书透明度和CA信息发现关联域名和服务")
    add_bullet(doc, "云资产发现：识别目标在云平台上的资产分布，包括对象存储、云主机、云数据库等")

    add_heading_text(doc, h5("供应链与第三方关联资产梳理"), 5)
    add_body(
        doc,
        "从攻击者视角来看，供应链和第三方资产往往是攻击突破的薄弱环节，需要重点关注："
    )
    add_bullet(doc, "供应商资产关联分析：梳理目标企业与供应商之间的网络连接、系统对接和数据交换关系")
    add_bullet(doc, "第三方服务平台：识别使用的第三方SaaS平台、CDN服务、API服务等")
    add_bullet(doc, "代码托管平台信息：在GitHub等代码托管平台搜索泄露的源码、配置文件和凭据信息")
    add_bullet(doc, "公共情报源收集：通过文库、网盘、论坛等公开渠道收集与目标相关的技术文档和敏感信息")

    add_heading_text(doc, h5("人员组织架构与社工信息收集"), 5)
    add_body(
        doc,
        "人员信息是社会工程学攻击的重要基础，通过合法的公开信息收集手段获取目标人员信息："
    )
    add_bullet(doc, "组织架构梳理：通过招聘网站、社交平台（LinkedIn等）收集目标单位组织架构和人员信息")
    add_bullet(doc, "邮箱账号收集：通过网络爬虫和搜索引擎语法发现目标单位邮箱账号和命名规则")
    add_bullet(doc, "社交媒体信息：收集目标人员在社交媒体上公开的信息，包括职务信息、联系方式等")
    add_bullet(doc, "密码泄露数据库排查：在合法的框架下排查目标邮箱是否曾出现在已知密码泄露事件中")

    add_heading_text(doc, h4("攻击战术制定"), 4)
    add_body(
        doc,
        "基于攻击面分析结果，结合MITRE ATT&CK攻击框架和实战经验，制定多阶段、多层次、多路径的攻击战术方案。"
    )

    add_heading_text(doc, h5("基于ATT&CK攻击框架的战术选型"), 5)
    add_body(
        doc,
        "依据MITRE ATT&CK框架，针对目标系统特性和技术栈，选择最优的攻击战术组合："
    )
    add_body(
        doc,
        "攻击战术的选型需综合考虑目标系统的技术架构、暴露面特征、安全防护措施以及攻击队的技术能力。"
        "通过将ATT&CK框架与实际目标特征进行映射，确保选型精准、覆盖全面。"
    )
    add_table_2col(
        doc,
        ["ATT&CK战术阶段", "对应攻击活动"],
        [
            ["初始访问（Initial Access）", "钓鱼攻击、Web漏洞利用、凭证爆破、供应链攻击"],
            ["执行（Execution）", "远控木马执行、Powershell脚本、恶意宏文件"],
            ["持久化（Persistence）", "创建后门账号、注册表自启动项、计划任务"],
            ["权限提升（Privilege Escalation）", "内核漏洞利用、SUID提权、令牌窃取"],
            ["防御规避（Defense Evasion）", "进程注入、WAF绕过、日志清除、代码混淆"],
            ["凭据访问（Credential Access）", "哈希抓取、键盘记录、凭证转储"],
            ["发现发现（Discovery）", "网络共享枚举、账户枚举、系统信息收集"],
            ["横向移动（Lateral Movement）", "Pass-the-Hash、远程服务利用、RDP跳转"],
            ["收集（Collection）", "敏感文件归档、屏幕捕获、输入捕获"],
            ["命令与控制（C2）", "DNS隧道、HTTPS隐蔽通信、Domain Fronting"],
            ["数据外传（Exfiltration）", "通过加密信道外传数据、分片传输"],
        ],
    )

    add_heading_text(doc, h5("多阶段攻击链路设计"), 5)
    add_body(
        doc,
        "攻击链路采用多层次、多入口、多路径设计，确保在某一攻击路径受阻时可灵活切换："
    )
    add_bold_lead(doc, "信息收集阶段（第1-2天）：",
                  "全面开展目标信息收集和攻击面分析，完成资产清单、漏洞初筛和攻击路径预评估。")
    add_bold_lead(doc, "漏洞利用阶段（第3-5天）：",
                  "针对发现的漏洞进行利用验证，优先利用高危漏洞和0-Day/N-Day漏洞获取初始权限。")
    add_bold_lead(doc, "权限维持与扩散阶段（第5-8天）：",
                  "在获取初始权限后，进行权限提升、横向移动和持续访问建立，扩大控制范围。")
    add_bold_lead(doc, "目标达成与总结阶段（第9-10天）：",
                  "完成攻击目标（如获取核心系统权限、敏感数据等），整理攻击成果，输出攻击技术复盘报告。")

    add_heading_text(doc, h5("社会工程学攻击战术"), 5)
    add_body(
        doc,
        "社会工程学攻击是红队攻击体系中的重要组成部分，通过分析目标人员的心理特点和行为习惯，"
        "设计和实施精准的社会工程学攻击方案："
    )
    add_bullet(doc, "鱼叉邮件攻击：根据目标人员职务和业务特点定制钓鱼邮件内容，诱导点击恶意链接或附件")
    add_bullet(doc, "电话钓鱼：模拟IT支持人员、供应商等角色，通过电话获取账号信息或其他敏感数据")
    add_bullet(doc, "物理渗透配合：在授权范围内进行门禁尾随等手段进入办公区域，实现内网接入")
    add_bullet(doc, "即时通讯钓鱼：通过微信、钉钉等办公即时通讯工具进行伪装身份钓鱼")

    add_heading_text(doc, h4("实时战术研判"), 4)
    add_body(
        doc,
        "实时战术研判是确保攻击活动持续高效推进的关键机制，通过建立7×24小时远程研判支撑体系，"
        "为攻击队提供持续的技术分析和决策支持。"
    )

    add_heading_text(doc, h5("7×24小时远程研判支撑机制"), 5)
    add_body(
        doc,
        "建立覆盖整个攻击活动周期的远程研判支撑机制："
    )
    add_bullet(doc, "实时通讯渠道：建立专用即时通讯群组和应急联系通道，确保信息传递及时畅通")
    add_bullet(doc, "远程协同平台：搭建安全可靠的远程协同研判环境，支持技术信息共享和战术讨论")
    add_bullet(doc, "技术专家值守：安排具备丰富攻防经验的技术专家轮班值守，随时响应攻击队的战术咨询和技术求助")
    add_bullet(doc, "信息共享机制：建立实时的攻击信息共享机制，包括新发现资产、漏洞信息、攻击路径等")

    add_heading_text(doc, h5("攻击路径动态调整策略"), 5)
    add_body(
        doc,
        "攻击过程中面临的不确定性因素较多，需要建立动态的路径调整机制："
    )
    add_body(
        doc,
        "当已选攻击路径受阻时，研判团队应立即启动备选方案，基于前期攻击面分析和情报收集成果，"
        "快速切换到替代攻击入口或战术手段。动态调整策略包括："
    )
    add_bullet(doc, "路径切换：当主攻击路径受阻时，快速切换到备选攻击路径")
    add_bullet(doc, "战术调整：根据目标安全响应情况，调整攻击手法和工具使用策略")
    add_bullet(doc, "目标扩展：在已确认的安全范围内，扩展攻击目标至更多相关系统")
    add_bullet(doc, "工具替换：当特定工具被识别或拦截时，快速替换为替代工具和技术")

    add_heading_text(doc, h5("每日《战术执行情况简报》"), 5)
    add_body(
        doc,
        "攻击活动期间，每日定时输出《战术执行情况简报》，内容涵盖："
    )
    add_bullet(doc, "当日攻击进展：各攻击路径的推进情况和阶段性成果")
    add_bullet(doc, "目标情报更新：新发现的资产、漏洞和可利用信息")
    add_bullet(doc, "战术效果评估：已采用战术的执行效果和有效性分析")
    add_bullet(doc, "次日攻击计划：次日攻击重点、战术调整和资源配置建议")
    add_bullet(doc, "异常情况记录：遇到的防御措施和安全告警响应情况")

    # ══════════════════════════════════════════════════════════════════
    # CHAPTER 3
    # ══════════════════════════════════════════════════════════════════
    add_heading_text(doc, h3("漏洞验证协同"), 3)

    add_body(
        doc,
        "漏洞验证协同是连接漏洞发现与攻击利用的关键桥梁，通过建立标准化的漏洞验证流程和多团队协同机制，"
        "确保每个漏洞的可利用性得到充分验证，支撑攻击推演和战术优化。"
    )

    add_heading_text(doc, h4("漏洞挖掘与发现"), 4)
    add_body(
        doc,
        "采用自动化工具与人工分析相结合的方式，对授权目标系统进行全面漏洞挖掘："
    )

    add_heading_text(doc, h5("Web应用漏洞深度挖掘"), 5)
    add_body(
        doc,
        "针对目标Web应用系统，开展深度漏洞挖掘工作，覆盖OWASP Top 10及常见Web安全漏洞："
    )
    add_bullet(doc, "SQL注入漏洞：识别和验证SQL注入漏洞，包括常规注入、盲注、延时注入和时间注入")
    add_bullet(doc, "跨站脚本（XSS）：检测反射型、存储型和DOM-based XSS漏洞，验证利用可行性")
    add_bullet(doc, "服务端请求伪造（SSRF）：发现SSRF漏洞，验证其内网探测和远程代码执行利用条件")
    add_bullet(doc, "远程代码执行（RCE）：检测因反序列化、模板注入、命令注入等导致的远程代码执行漏洞")
    add_bullet(doc, "文件操作漏洞：发现文件包含、文件上传绕过、任意文件读取等漏洞")
    add_bullet(doc, "逻辑漏洞：检测权限绕过、越权访问、业务逻辑缺陷等逻辑层安全漏洞")
    add_bullet(doc, "API安全漏洞：检测REST API和GraphQL中的身份认证缺陷、参数篡改和批量赋值漏洞")

    add_heading_text(doc, h5("主机与网络设备漏洞检测"), 5)
    add_body(
        doc,
        "对目标网络范围内的主机系统和网络设备进行漏洞扫描和检测："
    )
    add_bullet(doc, "操作系统漏洞：检测Windows、Linux等操作系统已知安全漏洞，包括未修补的CVE漏洞")
    add_bullet(doc, "中间件漏洞：检测Nginx、Apache、Tomcat、WebLogic等中间件的已知漏洞和配置缺陷")
    add_bullet(doc, "数据库漏洞：检测MySQL、Oracle、SQL Server、Redis、MongoDB等数据库安全配置和漏洞")
    add_bullet(doc, "网络设备漏洞：检测路由器、交换机、防火墙等网络设备的固件漏洞和配置隐患")
    add_bullet(doc, "容器和编排平台：检测Docker、Kubernetes等容器平台的配置缺陷和安全漏洞")

    add_heading_text(doc, h5("弱口令与配置缺陷检查"), 5)
    add_body(
        doc,
        "弱口令和配置缺陷是最常见的攻击突破口之一，进行系统化检查："
    )
    add_bullet(doc, "弱口令检测：对SSH、RDP、MySQL、FTP、SMTP等常见服务进行弱口令爆破测试")
    add_bullet(doc, "默认口令检查：确认各类系统设备是否仍使用出厂默认口令")
    add_bullet(doc, "配置缺陷检查：检查CORS配置、安全响应头、TLS配置、目录列表等配置问题")
    add_bullet(doc, "敏感信息泄露检查：检查Git信息泄露、目录遍历、备份文件、错误页面信息泄露等")

    add_heading_text(doc, h4("漏洞利用可行性验证"), 4)
    add_body(
        doc,
        "漏洞利用可行性验证是确保漏洞可利用性的关键步骤，通过PoC验证和利用条件分析，评估每项漏洞的实际攻击价值。"
    )

    add_heading_text(doc, h5("PoC开发与利用验证流程"), 5)
    add_body(
        doc,
        "建立标准化的PoC开发与验证流程，确保每个漏洞的可利用性得到充分验证："
    )
    add_bullet(doc, "漏洞复现验证：基于漏洞原理和公开PoC，在测试环境中复现漏洞触发条件")
    add_bullet(doc, "PoC定制开发：针对目标环境特性，定制开发适配的PoC验证脚本")
    add_bullet(doc, "利用条件分析：分析漏洞利用所需的前提条件、约束因素和前置依赖")
    add_bullet(doc, "利用成功率评估：综合考虑目标防护措施、漏洞特征和攻击者技能，评估利用成功率")
    add_bullet(doc, "证据固定：对漏洞利用过程和结果进行完整记录和证据固定，包括截图、流量包、时间戳等")

    add_heading_text(doc, h5("绕过WAF/IPS等安全防护策略"), 5)
    add_body(
        doc,
        "面对日益成熟的WAF、IPS、RASP等安全防护产品，需要建立专业的绕过策略："
    )
    add_bullet(doc, "编码绕过：使用URL编码、Unicode编码、双重重编码等方式绕过规则检测")
    add_bullet(doc, "分片绕过：将攻击载荷分片传输，利用协议解析差异绕过检测")
    add_bullet(doc, "混淆绕过：利用注释、大小写变异、语法特性混淆攻击载荷")
    add_bullet(doc, "协议降级绕过：利用TLS协议降级或HTTP协议特性绕过中间层检测")
    add_bullet(doc, "逻辑拆分绕过：将攻击逻辑拆分为多次正常请求组合，规避单一请求检测")

    add_heading_text(doc, h4("协同机制"), 4)
    add_body(
        doc,
        "建立高效的多团队协同机制，确保漏洞验证工作与攻击行动紧密配合，形成完整的攻防闭环。"
    )

    add_heading_text(doc, h5("漏洞报告技术内容评审"), 5)
    add_body(
        doc,
        "建立多维度的漏洞报告评审机制，确保每份漏洞报告的技术准确性和有效性："
    )
    add_bullet(doc, "技术准确性评审：验证漏洞描述、影响范围、利用条件和修复建议的准确性")
    add_bullet(doc, "证据完整性评审：确认报告中包含了完整的利用过程截图、请求/响应报文和时间轴记录")
    add_bullet(doc, "危害等级评定：结合资产价值、漏洞严重性和实际利用条件，综合评定危害等级")
    add_bullet(doc, "修复建议评审：确保修复建议具体可行，包含临时缓解和根本修复两个层面")

    add_heading_text(doc, h5("多团队协同漏洞验证流程"), 5)
    add_body(
        doc,
        "多团队协同是提升漏洞验证效率的关键，建立标准化的协同流程："
    )
    add_bullet(doc, "漏洞发现通知：漏洞发现团队即时将新漏洞信息通报至协同平台")
    add_bullet(doc, "验证任务分配：研判团队根据漏洞类型和复杂度，分配验证任务至对应技术专家")
    add_bullet(doc, "实时进展同步：验证过程中实时同步进展、技术难点和绕过方法")
    add_bullet(doc, "验证结果确认：验证完成后确认漏洞可利用性，更新漏洞状态至最终结论")
    add_bullet(doc, "知识共享沉淀：每个漏洞验证完成后，将技术要点和绕过经验沉淀至团队知识库")

    add_heading_text(doc, h5("漏洞成果汇总与优先级排序"), 5)
    add_body(
        doc,
        "对攻击活动中发现的所有漏洞进行分类汇总和优先级排序，为后续攻击决策提供依据："
    )
    add_body(
        doc,
        "漏洞优先级排序采用三维度评估模型：漏洞严重性（CVSS评分）、资产重要性（业务价值和数据敏感度）、"
        "利用可实现性（到达条件和前提约束）。综合三个维度得分排序，确定漏洞的优先利用顺序。"
    )

    # ══════════════════════════════════════════════════════════════════
    # CHAPTER 4
    # ══════════════════════════════════════════════════════════════════
    add_heading_text(doc, h3("攻击路径推演"), 3)

    add_body(
        doc,
        "攻击路径推演是基于攻击面分析和漏洞验证成果，对可能攻击路径进行系统性建模和推演的过程。"
        "通过攻击路径推演，可以识别最优攻击路径、预判攻击过程中的关键节点和潜在障碍，"
        "为战术执行提供科学决策依据。"
    )

    add_heading_text(doc, h4("攻击路径建模"), 4)
    add_body(
        doc,
        "采用攻击图（Attack Graph）方法，对目标网络环境进行攻击路径建模，将资产节点、漏洞信息和网络拓扑"
        "映射为系统化的攻击图模型，支持多维度攻击路径分析和计算。"
    )

    add_heading_text(doc, h5("基于攻击图的路径推演方法"), 5)
    add_body(
        doc,
        "攻击图模型将目标环境抽象为节点和边的集合："
    )
    add_bullet(doc, "节点（Node）：代表目标环境中的各类资产，包括主机、应用、网络设备、安全设备等")
    add_bullet(doc, "边（Edge）：代表资产间的网络连接关系、信任关系和访问路径")
    add_bullet(doc, "漏洞标注：在每个节点上标注已发现的安全漏洞和可利用条件")
    add_bullet(doc, "权限状态：标注攻击者当前对各节点的访问权限状态（无权限、普通用户、管理员）")
    add_body(
        doc,
        "攻击图生成后，可进行最短路径分析、关键节点识别和攻击路径可视化，为攻击决策提供直观的支撑。"
    )

    add_heading_text(doc, h5("多入口点攻击路径分析"), 5)
    add_body(
        doc,
        "在攻击面分析基础上，识别多个潜在攻击入口点，针对每个入口点开展独立路径分析："
    )
    add_bullet(doc, "外部Web应用入口：通过Web漏洞获取初始访问权限")
    add_bullet(doc, "VPN/远程接入入口：通过凭证爆破或会话劫持获取远程访问权限")
    add_bullet(doc, "邮件系统入口：通过钓鱼攻击获取内部用户凭证和初始访问")
    add_bullet(doc, "供应链入口：通过第三方服务接口或供应商系统间接访问目标网络")
    add_bullet(doc, "移动端入口：通过移动应用漏洞或恶意APK获取初始访问")

    add_heading_text(doc, h5("最短路径与最优路径计算"), 5)
    add_body(
        doc,
        "基于攻击图模型，结合图中节点权重和边约束，计算从入口点到目标节点的最短攻击路径和最优攻击路径："
    )
    add_bullet(doc, "最短路径：在攻击图中找到从入口到目标所需经过的最少节点数和步骤数")
    add_bullet(doc, "最优路径：综合考虑成功概率、检测风险、时间成本和资源消耗，选择综合性价比最高的路径")
    add_bullet(doc, "路径冗余设计：以应对防御措施阻断、路径失效等不确定性，设计多条备选路径")
    add_bullet(doc, "路径风险预评估：预判每条攻击路径在各阶段被检测到的可能性及应对策略")

    add_heading_text(doc, h4("攻击链还原分析"), 4)
    add_body(
        doc,
        "攻击链还原分析是在攻击行动推进过程中和攻击完成后，对已执行的攻击行为进行系统性还原和分析的过程。"
        "通过攻击链还原，可以全面审视攻击过程、评估攻击成效、识别关键成功因素和改进空间。"
    )

    add_heading_text(doc, h5("攻击链全生命周期还原"), 5)
    add_body(
        doc,
        "参照业界通用的Cyber Kill Chain模型，对攻击行为进行全生命周期分阶段还原："
    )
    add_bullet(doc, "侦察跟踪（Reconnaissance）：还原攻击前期的信息收集和情报获取活动")
    add_bullet(doc, "武器构建（Weaponization）：还原漏洞利用工具、恶意载荷和钓鱼邮件的制作过程")
    add_bullet(doc, "载荷投递（Delivery）：还原攻击载荷投递方式和途径")
    add_bullet(doc, "漏洞利用（Exploitation）：还原漏洞利用的具体技术细节和触发条件")
    add_bullet(doc, "安装植入（Installation）：还原后门程序、远控木马的安装和启动过程")
    add_bullet(doc, "命令与控制（C2）：还原C2通信的建立、协议特征和流量隐蔽方式")
    add_bullet(doc, "目标达成（Actions on Objectives）：还原最终攻击目标的达成过程和成果获取情况")

    add_heading_text(doc, h5("关键攻击链节点识别"), 5)
    add_body(
        doc,
        "在攻击链还原过程中，识别对攻击成败具有决定性影响的关键节点，为后续攻击战术优化提供参考："
    )
    add_bullet(doc, "关键漏洞节点：对攻击链推进起到关键突破作用的漏洞")
    add_bullet(doc, "关键权限节点：攻击链中关键权限提升和横向移动的操作节点")
    add_bullet(doc, "关键决策节点：攻击过程中对路径选择和技术方案做出关键决策的节点")
    add_bullet(doc, "风险暴露节点：攻击过程中可能存在暴露风险的操作节点和对应检测点")

    add_heading_text(doc, h5("攻击路径可视化呈现"), 5)
    add_body(
        doc,
        "采用可视化的方式呈现攻击路径和攻击链还原结果，使技术团队和管理层能够直观理解攻击过程："
    )
    add_bullet(doc, "网络拓扑图标注：在目标网络拓扑图上标注攻击入口、渗透路径和控制节点")
    add_bullet(doc, "攻击链时序图：以时间轴方式呈现攻击链各阶段的关键事件和操作节点")
    add_bullet(doc, "权限变化图：呈现攻击过程中攻击者在各节点上权限状态的变化过程")
    add_bullet(doc, "攻击流量流向图：呈现攻击流量在网络中的传播路径和跳转关系")

    add_heading_text(doc, h4("场景化推演"), 4)
    add_body(
        doc,
        "针对不同的攻击场景，开展差异化的攻击路径推演，确保推演结果符合各场景的实际特点和需求。"
    )

    add_heading_text(doc, h5("网安行动场景攻防推演"), 5)
    add_body(
        doc,
        "网安行动场景下的攻击路径推演重点关注：攻击的隐蔽性、持续性和深度渗透能力。"
        "推演将模拟攻击者从外部互联网入口逐步渗透至内部核心网络的全过程，"
        "重点关注安全防御体系的盲区和薄弱环节。"
    )

    add_heading_text(doc, h5("众测场景深度渗透推演"), 5)
    add_body(
        doc,
        "众测场景下的攻击路径推演侧重于：漏洞的全面发现和利用验证有效性。"
        "推演将以单点突破为核心，围绕单个系统或多个关联系统开展深度漏洞挖掘和利用验证，"
        "确保在有限时间内最大化漏洞发现效率。"
    )

    add_heading_text(doc, h5("竞赛场景技术突破推演"), 5)
    add_body(
        doc,
        "竞赛场景下的攻击路径推演侧重于：速度、准确性和技术广度。"
        "推演将针对竞赛可能涉及的常见技术栈和漏洞类型，设计快速突破的攻击路径模板，"
        "并在赛前通过集训验证和优化推演成果。"
    )

    # ══════════════════════════════════════════════════════════════════
    # CHAPTER 5
    # ══════════════════════════════════════════════════════════════════
    add_heading_text(doc, h3("复盘分析"), 3)

    add_body(
        doc,
        "复盘分析是攻击活动闭环管理的关键环节，通过对攻击全过程、全要素的系统性回顾和分析，"
        "总结经验教训、提炼最佳实践、形成可复用的攻击知识库，持续提升红队攻击能力。"
    )

    add_heading_text(doc, h4("攻击技术复盘"), 4)
    add_body(
        doc,
        "攻击技术复盘是整个复盘分析的核心，从技术层面对攻击活动进行全面回顾和评估。"
    )

    add_heading_text(doc, h5("漏洞利用技术复盘"), 5)
    add_body(
        doc,
        "对攻击活动中涉及的所有漏洞利用技术进行系统复盘："
    )
    add_bullet(doc, "技术手段回顾：回顾每个漏洞利用过程中采用的技术手段、工具和技巧")
    add_bullet(doc, "利用成功率统计：统计各类漏洞利用技术的成功率和效率指标")
    add_bullet(doc, "技术难点分析：分析利用过程中遇到的技术障碍和绕过技术")
    add_bullet(doc, "创新技术总结：总结本次攻击活动中开发或应用的技术创新点")
    add_bullet(doc, "技术能力差距：识别团队在技术能力上存在的差距和改进方向")

    add_heading_text(doc, h5("攻击路径有效性复盘"), 5)
    add_body(
        doc,
        "对攻击活动中执行的所有攻击路径进行有效性评估："
    )
    add_bullet(doc, "路径执行效果：评估每条攻击路径的执行效果，包括用时、成功率、隐蔽性等指标")
    add_bullet(doc, "路径选择合理性：回顾路径选择过程，评估路径选择决策的合理性")
    add_bullet(doc, "路径切换效果：评估攻击过程中路径切换的及时性和有效性")
    add_bullet(doc, "路径优化建议：基于复盘分析结果，提出攻击路径设计的优化建议")

    add_heading_text(doc, h5("战术执行效果复盘"), 5)
    add_body(
        doc,
        "对攻击活动中采用的各种战术进行执行效果评估："
    )
    add_bullet(doc, "战术命中率：评估各战术在实际攻击中的命中成功率和贡献度")
    add_bullet(doc, "战术适应性：分析各类战术在不同目标环境下的适应性和局限性")
    add_bullet(doc, "战术组合效果：评估多战术组合使用的协同效果和互相支持情况")
    add_bullet(doc, "战术改进建议：基于复盘结果，提出战术选型和执行的改进建议")

    add_heading_text(doc, h4("攻击链还原与根因分析"), 4)
    add_body(
        doc,
        "在攻击技术复盘的基础上，对关键攻击链进行深度还原和根因分析："
    )

    add_heading_text(doc, h5("关键攻击链还原记录"), 5)
    add_body(
        doc,
        "对攻击活动中最具代表性的关键攻击链进行完整还原记录："
    )
    add_bullet(doc, "时间轴还原：以精确到小时的时间粒度还原攻击链的完整过程")
    add_bullet(doc, "技术细节记录：记录攻击链每个环节的技术操作细节、工具版本和命令执行序列")
    add_bullet(doc, "决策点记录：记录攻击链中的关键决策点、决策依据和替代方案")
    add_bullet(doc, "异常事件记录：记录攻击过程中遇到的异常事件、应对措施和效果")

    add_heading_text(doc, h5("漏洞根因分析与归类"), 5)
    add_body(
        doc,
        "对攻击活动中发现的漏洞进行根因分析和归类："
    )
    add_bullet(doc, "根因分析：分析每个漏洞的根根本原因（开发缺陷、配置疏漏、运维疏忽等）")
    add_bullet(doc, "漏洞归类：将漏洞按照OWASP、CWE等标准进行分类统计")
    add_bullet(doc, "分布特征分析：分析漏洞在系统类型、部署位置、开发语言等维度上的分布特征")
    add_bullet(doc, "趋势分析：对比历次攻击活动的漏洞发现数据，分析漏洞趋势变化")

    add_heading_text(doc, h5("攻击成功率与效率统计分析"), 5)
    add_body(
        doc,
        "对攻击活动的成功率和效率进行量化统计分析："
    )
    add_body(
        doc,
        "统计指标覆盖攻击成功率（已成功利用漏洞数/发现漏洞总数）、攻击效率（发现一个可利用漏洞的平均用时）、"
        "目标覆盖率（成功渗透的目标数/目标总数）和攻击路径有效率（有效路径数/总路径数）等核心指标。"
        "通过多维度的量化分析，客观评估攻击活动执行效果和团队能力表现。"
    )

    add_heading_text(doc, h4("《攻击技术复盘报告》输出"), 4)
    add_body(
        doc,
        "攻击活动结束后，输出完整的《攻击技术复盘报告》，作为攻击活动的总结性交付物。"
    )

    add_heading_text(doc, h5("报告结构规范"), 5)
    add_body(
        doc,
        "《攻击技术复盘报告》采用标准化结构，确保内容全面、结构清晰："
    )
    add_bullet(doc, "活动概述：攻击活动基本信息（时间、范围、目标、团队构成）")
    add_bullet(doc, "攻击过程记录：攻击全过程时序记录，包括关键事件和决策节点")
    add_bullet(doc, "漏洞发现清单：所有发现漏洞的详细清单，含危害等级和利用情况")
    add_bullet(doc, "攻击链还原：关键攻击链的完整还原记录和可视化呈现")
    add_bullet(doc, "成果统计分析：攻击成果的量化统计和多维分析")
    add_bullet(doc, "经验与教训：攻击过程中的成功经验和值得改进的方面")
    add_bullet(doc, "防御改进建议：从攻击者视角提出的防守侧补防建议")

    add_heading_text(doc, h5("关键发现与改进建议"), 5)
    add_body(
        doc,
        "基于攻击活动的执行情况和复盘分析结果，提炼关键发现和系统化的改进建议："
    )
    add_body(
        doc,
        "改进建议分为三个层次：战术层（攻击战术和技术的改进方向）、能力层（团队技术能力的提升方向）"
        "和工具层（攻击工具和基础设施的改进方向），形成系统化、可落地的改进路线图。"
    )

    add_heading_text(doc, h5("防守侧补防建议"), 5)
    add_body(
        doc,
        "从攻击者视角出发，针对攻击活动中暴露的安全短板和防御盲区，提出防守侧的具体补防建议："
    )
    add_bullet(doc, "暴露面收敛：建议关闭/收敛的互联网暴露端口、服务和应用系统")
    add_bullet(doc, "漏洞修复优先级：基于实际可利用性提出的漏洞修复优先级排序")
    add_bullet(doc, "检测规则优化：针对防守侧检测盲区提出的检测规则优化建议")
    add_bullet(doc, "应急响应改进：针对防守侧应急响应过程中的不足提出的改进建议")
    add_bullet(doc, "安全培训建议：基于攻击过程中发现的人员安全意识问题提出的培训建议")

    # ══════════════════════════════════════════════════════════════════
    # CHAPTER 6
    # ══════════════════════════════════════════════════════════════════
    add_heading_text(doc, h3("服务团队与保障"), 3)

    add_body(
        doc,
        "为确保红队专项攻击支撑服务的高质量交付，组建具备丰富实战经验的红队专家团队，"
        "并建立完善的服务保障体系。"
    )

    add_heading_text(doc, h4("红队专家团队配置"), 4)
    add_body(
        doc,
        "项目团队由多名具备丰富红队攻击实战经验的技术专家组成，核心成员不少于3人，涵盖以下角色："
    )

    add_table_2col(
        doc,
        ["角色", "职责描述"],
        [
            ["项目经理", "负责项目整体管理、资源协调、质量控制和客户沟通，统领攻击研判技术方向"],
            ["高级攻击研究员", "负责攻击战术研判、攻击路径推演和PoC开发，指导关键攻击技术的攻克"],
            ["渗透测试工程师", "负责漏洞挖掘、利用验证和攻击执行，执行日常渗透测试任务"],
            ["情报分析员", "负责攻击面分析、目标情报收集和社会工程学攻击支撑"],
            ["复盘分析师", "负责攻击过程记录、攻击链还原和复盘分析报告编制"],
        ],
    )

    add_heading_text(doc, h4("人员资质要求"), 4)
    add_body(
        doc,
        "项目团队成员具备以下资质要求："
    )
    add_bullet(doc, "项目经理：本科及以上学历，5年以上网络信息安全从业经验；具备CISP（注册信息安全专业人员）或CCRC-DSO认证资质")
    add_bullet(doc, "安全工程师：3年以上网络安全从业经验；具备CISP、CISAW（信息安全保障人员认证证书）等安全资质")
    add_bullet(doc, "熟练掌握各类安全扫描工具和渗透测试工具的使用，包括Burp Suite、Nmap、Metasploit、Cobalt Strike等")
    add_bullet(doc, "具备丰富的红队攻击实战经验，参与过省市级或国家级攻防演练活动")

    add_heading_text(doc, h4("服务承诺与质量保障"), 4)

    add_heading_text(doc, h5("合法合规承诺"), 5)
    add_body(
        doc,
        "应答人郑重承诺在服务过程中严格遵守《网络安全法》及相关法律法规，所有红队攻击行为仅限于授权目标范围，"
        "不得触及非授权系统。所有攻击活动均在采购人明确的书面授权范围内执行，"
        "严格遵循授权边界进行操作，确保红队攻击活动的合法合规性。"
    )

    add_heading_text(doc, h5("保密协议"), 5)
    add_body(
        doc,
        "参加此次项目的所有项目组成员与采购人签署相关的保密协议和非侵害协议。"
        "对在服务过程中获取的采购人系统信息、漏洞信息、敏感数据等严格保密，"
        "未经采购人书面许可，不得向任何第三方披露。服务结束后，全部技术资料和数据移交采购人或按规定销毁。"
    )

    add_heading_text(doc, h5("服务质量考核标准"), 5)
    add_body(
        doc,
        "服务质量考核依据采购人考核标准和细则执行，主要考核维度包括："
    )
    add_bullet(doc, "攻击成效指标（50%）：包括漏洞发现数量和质量、攻击得分水平、目标达成率等")
    add_bullet(doc, "技术能力表现（25%）：包括技术手段的先进性和多样性、技术难点的攻克能力")
    add_bullet(doc, "服务交付质量（15%）：包括报告的时效性、准确性和完整性，战术研判的及时性和有效性")
    add_bullet(doc, "合规安全表现（10%）：包括授权遵守、操作规范和保密执行情况")

    # ── save ─────────────────────────────────────────────────────────
    doc.save(OUTPUT)
    print(f"Saved to {OUTPUT}")


if __name__ == "__main__":
    doc = Document(str(TEMPLATE))
    build(doc)
