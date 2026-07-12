from __future__ import annotations

import argparse
import re
import tempfile
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


BODY_FIRST_LINE_INDENT = Pt(28)
BODY_FONT_SIZE = Pt(11)
TABLE_FONT_SIZE = Pt(9)


def clear_document(document: Document) -> None:
    body = document._element.body
    for child in list(body):
        if child.tag.endswith("sectPr"):
            continue
        body.remove(child)


def set_run_font(run, size=BODY_FONT_SIZE, bold=None):
    run.font.size = size
    if bold is not None:
        run.bold = bold
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")


def set_paragraph_font(paragraph, size=BODY_FONT_SIZE, bold=None):
    for run in paragraph.runs:
        set_run_font(run, size=size, bold=bold)


def style_body_paragraph(paragraph, first_line=True):
    paragraph.paragraph_format.first_line_indent = BODY_FIRST_LINE_INDENT if first_line else None
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.28
    set_paragraph_font(paragraph)


def first_existing_style(document: Document, names):
    for name in names:
        try:
            document.styles[name]
            return name
        except KeyError:
            continue
    return ""


def iter_block_items(document: Document):
    body = document.element.body
    for child in body.iterchildren():
        if child.tag == qn("w:p"):
            yield "paragraph", child
        elif child.tag == qn("w:tbl"):
            yield "table", child


def paragraph_by_element(document: Document, element):
    for paragraph in document.paragraphs:
        if paragraph._element is element:
            return paragraph
    return None


def table_by_element(document: Document, element):
    for table in document.tables:
        if table._element is element:
            return table
    return None


def clean_text(text: str) -> str:
    text = text.replace("\u00a0", " ")
    return re.sub(r"[ \t]+", " ", text).strip()


def strip_source_bullet(text: str) -> tuple[bool, str]:
    stripped = clean_text(text)
    bullet_match = re.match(r"^[●•Øo]\s*(.+)$", stripped)
    if bullet_match:
        return True, clean_text(bullet_match.group(1))
    return False, stripped


def source_heading_level(paragraph, text: str, previous_heading: int, source_min_level: int = 1, start_level: int = 1) -> int:
    style_name = paragraph.style.name if paragraph.style else ""
    if style_name.startswith("Heading") or style_name.startswith("标题"):
        match = re.search(r"(\d+)", style_name)
        raw_level = int(match.group(1)) if match else 4
        return min(8, max(start_level, raw_level - source_min_level + start_level))
    if style_name in ("五级",):
        return min(8, max(start_level, 5 - source_min_level + start_level))
    if style_name in ("六级",):
        return min(8, max(start_level, 6 - source_min_level + start_level))
    return 0


def image_rids(paragraph):
    rids = []
    for node in paragraph._element.iter():
        if node.tag.endswith("}blip"):
            rid = node.get(qn("r:embed"))
            if rid:
                rids.append(rid)
    return rids


def add_heading(document: Document, text: str, level: int):
    paragraph = document.add_heading(text, level=level)
    paragraph.paragraph_format.space_before = Pt(12 if level == 1 else 8)
    paragraph.paragraph_format.space_after = Pt(6)
    set_paragraph_font(paragraph, size=Pt(16 if level == 1 else 14 if level == 2 else 12), bold=True)
    return paragraph


def add_bullet(document: Document, text: str):
    style = first_existing_style(document, ["List Bullet", "List Paragraph", "列表段落"])
    if style and style != "List Paragraph":
        paragraph = document.add_paragraph(style=style)
        paragraph.add_run(text)
    else:
        paragraph = document.add_paragraph(f"• {text}")
        paragraph.paragraph_format.left_indent = Pt(18)
        paragraph.paragraph_format.first_line_indent = None
    paragraph.paragraph_format.space_after = Pt(3)
    set_paragraph_font(paragraph)


def add_body(document: Document, text: str):
    paragraph = document.add_paragraph(text)
    style_body_paragraph(paragraph)


def set_cell_text(cell, text: str, header=False):
    cell.text = clean_text(text)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for paragraph in cell.paragraphs:
        paragraph.paragraph_format.first_line_indent = None
        paragraph.paragraph_format.space_after = Pt(2)
        paragraph.paragraph_format.line_spacing = 1.15
        if header:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in paragraph.runs:
            set_run_font(run, size=TABLE_FONT_SIZE, bold=header)


def shade_cell(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def copy_table(document: Document, source_table):
    if not source_table.rows:
        return
    rows = len(source_table.rows)
    cols = max(len(row.cells) for row in source_table.rows)
    table = document.add_table(rows=rows, cols=cols)
    table.style = "Table Grid"
    table.autofit = True
    for row_index, row in enumerate(source_table.rows):
        for col_index, source_cell in enumerate(row.cells):
            if col_index >= len(table.rows[row_index].cells):
                continue
            cell = table.rows[row_index].cells[col_index]
            set_cell_text(cell, source_cell.text, header=row_index == 0)
            if row_index == 0:
                shade_cell(cell, "D9EAF7")
    document.add_paragraph("")


def add_paragraph_images(document: Document, source_doc: Document, paragraph):
    for rid in image_rids(paragraph):
        part = source_doc.part.related_parts.get(rid)
        if not part:
            continue
        suffix = Path(part.partname).suffix or ".png"
        with tempfile.NamedTemporaryFile(suffix=suffix, delete=False) as tmp:
            tmp.write(part.blob)
            tmp_path = Path(tmp.name)
        try:
            document.add_picture(str(tmp_path), width=Inches(5.6))
            last = document.paragraphs[-1]
            last.alignment = WD_ALIGN_PARAGRAPH.CENTER
            last.paragraph_format.space_after = Pt(8)
        finally:
            tmp_path.unlink(missing_ok=True)


def source_base_heading_level(document: Document) -> int:
    for paragraph in document.paragraphs:
        style_name = paragraph.style.name if paragraph.style else ""
        if style_name.startswith("Heading") or style_name.startswith("标题"):
            match = re.search(r"(\d+)", style_name)
            if match and paragraph.text.strip():
                return int(match.group(1))
    return 1


def convert(source_path: Path, template_path: Path, output_path: Path, start_level: int = 1, include_title: bool = True):
    source = Document(source_path)
    target = Document(template_path)
    clear_document(target)
    source_min_level = source_base_heading_level(source)

    for section in target.sections:
        section.top_margin = Inches(0.9)
        section.bottom_margin = Inches(0.9)
        section.left_margin = Inches(0.9)
        section.right_margin = Inches(0.9)

    last_heading = start_level
    if include_title:
        add_heading(target, source_path.stem, start_level)

    for kind, element in iter_block_items(source):
        if kind == "paragraph":
            paragraph = paragraph_by_element(source, element)
            if paragraph is None:
                continue
            raw_text = paragraph.text
            is_bullet, text = strip_source_bullet(raw_text)
            if not text:
                add_paragraph_images(target, source, paragraph)
                continue
            if text in {"测试目的", "测试方法"}:
                p = target.add_paragraph(text)
                p.paragraph_format.first_line_indent = None
                p.paragraph_format.space_before = Pt(4)
                p.paragraph_format.space_after = Pt(2)
                set_paragraph_font(p, bold=True)
            elif is_bullet:
                add_bullet(target, text)
            else:
                level = source_heading_level(paragraph, text, last_heading, source_min_level, start_level)
                if level:
                    add_heading(target, text, level)
                    last_heading = level
                elif paragraph.style.name in ("列表段落1",):
                    p = target.add_paragraph(text)
                    p.paragraph_format.first_line_indent = None
                    p.paragraph_format.space_after = Pt(3)
                    p.paragraph_format.space_before = Pt(4)
                    set_paragraph_font(p, bold=True)
                elif paragraph.style.name in ("长亭正文",):
                    p = target.add_paragraph(text)
                    p.paragraph_format.first_line_indent = BODY_FIRST_LINE_INDENT
                    p.paragraph_format.space_after = Pt(6)
                    p.paragraph_format.line_spacing = 1.28
                    set_paragraph_font(p, bold=True)
                else:
                    add_body(target, text)
            add_paragraph_images(target, source, paragraph)
        elif kind == "table":
            table = table_by_element(source, element)
            if table is not None:
                copy_table(target, table)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    target.save(output_path)


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("source")
    parser.add_argument("--template", default="templates/default-proposal-template.docx")
    parser.add_argument("--out", required=True)
    parser.add_argument("--start-level", type=int, default=1)
    parser.add_argument("--no-title", action="store_true")
    args = parser.parse_args()
    convert(Path(args.source), Path(args.template), Path(args.out), start_level=args.start_level, include_title=not args.no_title)


if __name__ == "__main__":
    main()
