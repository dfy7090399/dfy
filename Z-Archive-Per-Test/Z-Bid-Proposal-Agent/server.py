from __future__ import annotations

import cgi
import ipaddress
import json
import logging
import os
import re
import shutil
import socket
import subprocess
import threading
import time
import urllib.error
import urllib.request
import uuid
from http.cookies import SimpleCookie
from datetime import datetime, timezone
from http import HTTPStatus
from http.server import SimpleHTTPRequestHandler, ThreadingHTTPServer
from pathlib import Path
from urllib.parse import unquote, urlsplit

from docx import Document
from docx.shared import Pt

from agents import run_pipeline


BASE_DIR = Path(__file__).resolve().parent
SRC_DIR = BASE_DIR / "src"
TEMPLATE_DIR = BASE_DIR / "templates"
OUTPUT_DIR = BASE_DIR / "output"
CONFIG_DIR = BASE_DIR / "config"
MODEL_CONFIG_FILE = CONFIG_DIR / "model-config.json"
MODEL_CONFIG_LOCAL_FILE = CONFIG_DIR / "model-config.local.json"
DEFAULT_TEMPLATE = TEMPLATE_DIR / "default-proposal-template.docx"
DEFAULT_STYLE = TEMPLATE_DIR / "default-template-style.json"
TARGET_CHARS_PER_PAGE = 590
MIN_CHILD_CHARS = 260
BODY_FIRST_LINE_INDENT = Pt(28)
COMPOSE_STATUS_FILE = "compose-status.json"
ACCESS_CONTROL_FILE = ".access.json"
ACTIVE_COMPOSE_TASKS: set[str] = set()
ACTIVE_COMPOSE_LOCK = threading.Lock()
SESSION_COOKIE_NAME = "bid_proposal_session"
MAX_JSON_BODY_BYTES = 256 * 1024
MAX_UPLOAD_BYTES = 10 * 1024 * 1024
MAX_MULTIPART_BYTES = 25 * 1024 * 1024
MAX_DOCX_XML_BYTES = 10 * 1024 * 1024
ALLOWED_OPENAI_HOSTS = {"api.openai.com"}
ALLOWED_OPENAI_SUFFIXES = (".openai.azure.com", ".services.ai.azure.com")
ALLOWED_ANTHROPIC_HOSTS = {"api.anthropic.com"}
DISALLOWED_FILE_SEGMENTS = {"uploads", COMPOSE_STATUS_FILE, ACCESS_CONTROL_FILE}
ALLOWED_UPLOAD_SUFFIXES = {
    "technicalSpecification": {".txt", ".md", ".csv", ".docx", ".xlsx", ".xls", ".json"},
    "scoringCriteria": {".txt", ".md", ".csv", ".docx", ".xlsx", ".xls", ".json"},
    "templateFile": {".docx"},
}
RUNTIME_SECRETS = {
    "anthropic": "",
    "openai": "",
}

DEFAULT_MODEL_CONFIG = {
    "writer": {
        "providerPriority": ["codex", "openai", "local"],
    },
    "codex": {
        "enabled": True,
        "useCurrentLogin": True,
        "model": "",
        "timeoutSeconds": 240,
    },
    "openai": {
        "enabled": True,
        "baseUrl": "https://api.openai.com/v1",
        "model": "gpt-4o-mini",
        "apiKeyEnv": "OPENAI_API_KEY",
        "apiKey": "",
    },
    "local": {
        "enabled": True,
    },
}


def deep_merge(base, override):
    merged = dict(base)
    for key, value in override.items():
        if isinstance(value, dict) and isinstance(merged.get(key), dict):
            merged[key] = deep_merge(merged[key], value)
        else:
            merged[key] = value
    return merged


def read_json_file(path):
    if not path.exists():
        return {}
    try:
        return json.loads(path.read_text(encoding="utf-8"))
    except json.JSONDecodeError:
        return {}


def load_model_config():
    config = deep_merge(DEFAULT_MODEL_CONFIG, read_json_file(MODEL_CONFIG_FILE))
    config = deep_merge(config, read_json_file(MODEL_CONFIG_LOCAL_FILE))
    config_env = config.get("env", {})
    if config_env:
        config = deep_merge(config, {
            "writer": {
                "providerPriority": ["anthropic", "codex", "openai", "local"],
            },
            "anthropic": {
                "enabled": True,
                "baseUrl": config_env.get("ANTHROPIC_BASE_URL", ""),
                "model": config_env.get("ANTHROPIC_MODEL")
                    or config_env.get("ANTHROPIC_DEFAULT_SONNET_MODEL")
                    or config_env.get("ANTHROPIC_DEFAULT_OPUS_MODEL")
                    or config_env.get("ANTHROPIC_DEFAULT_HAIKU_MODEL")
                    or "",
                "apiKeyEnv": "ANTHROPIC_AUTH_TOKEN",
                "apiKey": config_env.get("ANTHROPIC_AUTH_TOKEN", ""),
            },
        })
    return config


def env_bool(name, default):
    value = os.environ.get(name)
    if value is None:
        return default
    return value.strip().lower() not in {"0", "false", "no", "off"}


class ProposalHandler(SimpleHTTPRequestHandler):
    def __init__(self, *args, **kwargs):
        self._session_cookie_value = None
        super().__init__(*args, directory=str(SRC_DIR), **kwargs)

    def end_headers(self):
        if self._session_cookie_value:
            self.send_header(
                "Set-Cookie",
                f"{SESSION_COOKIE_NAME}={self._session_cookie_value}; Path=/; HttpOnly; SameSite=Strict",
            )
            self._session_cookie_value = None
        super().end_headers()

    def current_session_id(self):
        cookie_header = self.headers.get("Cookie", "")
        if not cookie_header:
            return ""
        cookie = SimpleCookie()
        try:
            cookie.load(cookie_header)
        except Exception:
            return ""
        morsel = cookie.get(SESSION_COOKIE_NAME)
        if not morsel:
            return ""
        return morsel.value.strip()

    def ensure_session(self):
        session_id = self.current_session_id()
        if session_id:
            return session_id
        session_id = uuid.uuid4().hex
        self._session_cookie_value = session_id
        return session_id

    def require_session(self):
        session_id = self.current_session_id()
        if not session_id:
            raise PermissionError("当前会话未授权，请先刷新页面后重试")
        return session_id

    def do_POST(self):
        try:
            self.require_session()
            if self.path == "/api/model-config":
                response = self.handle_model_config_save()
                self.send_json(response, HTTPStatus.OK)
                return

            if self.path == "/api/model-config/test":
                response = self.handle_model_config_test()
                self.send_json(response, HTTPStatus.OK)
                return

            if self.path == "/api/proposals":
                response = self.handle_proposal_request()
                self.send_json(response, HTTPStatus.CREATED)
                return

            if self.path.startswith("/api/proposals/") and self.path.endswith("/outline/chat"):
                job_id = self.path.removeprefix("/api/proposals/").removesuffix("/outline/chat").strip("/")
                response = self.handle_outline_chat(job_id)
                self.send_json(response, HTTPStatus.OK)
                return

            if self.path.startswith("/api/proposals/") and self.path.endswith("/compose"):
                job_id = self.path.removeprefix("/api/proposals/").removesuffix("/compose").strip("/")
                response = self.handle_compose_request(job_id)
                self.send_json(response, HTTPStatus.ACCEPTED)
                return

            self.send_error(HTTPStatus.NOT_FOUND, "Unknown API endpoint")
            return
        except ValueError as exc:
            self.send_json({"error": str(exc)}, HTTPStatus.BAD_REQUEST)
            return
        except PermissionError as exc:
            self.send_json({"error": str(exc)}, HTTPStatus.FORBIDDEN)
            return
        except Exception as exc:  # pragma: no cover - local server guardrail
            logging.exception("Unhandled POST error")
            self.send_json({"error": f"处理失败：{exc}"}, HTTPStatus.INTERNAL_SERVER_ERROR)
            return

    def do_GET(self):
        if not self.path.startswith("/api/") and not self.path.startswith("/output/"):
            self.ensure_session()

        if self.path == "/api/model-config":
            try:
                self.require_session()
                self.send_json(self.handle_model_config_read(), HTTPStatus.OK)
            except PermissionError as exc:
                self.send_json({"error": str(exc)}, HTTPStatus.FORBIDDEN)
            return

        if self.path == "/api/proposals/latest":
            try:
                self.require_session()
                self.send_json(self.handle_latest_proposal(), HTTPStatus.OK)
            except ValueError as exc:
                self.send_json({"error": str(exc)}, HTTPStatus.NOT_FOUND)
            except PermissionError as exc:
                self.send_json({"error": str(exc)}, HTTPStatus.FORBIDDEN)
            return

        if self.path.startswith("/api/proposals/") and self.path.endswith("/compose/status"):
            try:
                self.require_session()
                job_id = self.path.removeprefix("/api/proposals/").removesuffix("/compose/status").strip("/")
                self.send_json(self.handle_compose_status(job_id), HTTPStatus.OK)
            except ValueError as exc:
                self.send_json({"error": str(exc)}, HTTPStatus.BAD_REQUEST)
            except PermissionError as exc:
                self.send_json({"error": str(exc)}, HTTPStatus.FORBIDDEN)
            return

        if self.path.startswith("/api/proposals/"):
            try:
                self.require_session()
                job_id = self.path.removeprefix("/api/proposals/").strip("/")
                self.send_json(self.handle_proposal_detail(job_id), HTTPStatus.OK)
            except ValueError as exc:
                self.send_json({"error": str(exc)}, HTTPStatus.NOT_FOUND)
            except PermissionError as exc:
                self.send_json({"error": str(exc)}, HTTPStatus.FORBIDDEN)
            return

        if self.path.startswith("/output/"):
            try:
                self.require_session()
                self.serve_output_file()
            except PermissionError as exc:
                self.send_error(HTTPStatus.FORBIDDEN, str(exc))
            return

        super().do_GET()

    def do_HEAD(self):
        if self.path.startswith("/output/"):
            try:
                self.require_session()
                self.serve_output_file()
            except PermissionError as exc:
                self.send_error(HTTPStatus.FORBIDDEN, str(exc))
            return

        self.ensure_session()
        super().do_HEAD()

    def handle_model_config_read(self):
        config = self.config_for_display()
        return {
            "config": self.mask_model_config(config),
            "activeProvider": self.active_writer_provider(),
            "engine": self.writer_engine_name(),
            "model": self.writer_model_name(),
            "localConfigExists": MODEL_CONFIG_LOCAL_FILE.exists(),
        }

    def handle_model_config_save(self):
        payload = self.parse_json_body()
        provider = payload.get("provider", "anthropic")
        current = load_model_config()
        current_env = current.get("env", {})
        current_openai = current.get("openai", {})
        anthropic_key = payload.get("anthropicApiKey", "")
        openai_key = payload.get("openaiApiKey", "")
        if not anthropic_key or anthropic_key == "********":
            anthropic_key = RUNTIME_SECRETS["anthropic"] or current_env.get("ANTHROPIC_AUTH_TOKEN", "")
        if not openai_key or openai_key == "********":
            openai_key = RUNTIME_SECRETS["openai"] or current_openai.get("apiKey", "")
        anthropic_base_url = self.validate_base_url(payload.get("anthropicBaseUrl", ""), "anthropic", allow_blank=True)
        openai_base_url = self.validate_base_url(payload.get("openaiBaseUrl", "https://api.openai.com/v1"), "openai")
        RUNTIME_SECRETS["anthropic"] = anthropic_key
        RUNTIME_SECRETS["openai"] = openai_key
        config = {
            "env": {
                "ANTHROPIC_AUTH_TOKEN": "",
                "ANTHROPIC_BASE_URL": anthropic_base_url,
                "ANTHROPIC_MODEL": payload.get("anthropicModel", ""),
            },
            "writer": {
                "providerPriority": payload.get("providerPriority") or self.provider_priority_for(provider),
            },
            "codex": {
                "enabled": bool(payload.get("codexEnabled", True)),
                "model": payload.get("codexModel", ""),
                "timeoutSeconds": int(payload.get("codexTimeoutSeconds") or 240),
            },
            "openai": {
                "enabled": bool(payload.get("openaiEnabled", True)),
                "baseUrl": openai_base_url,
                "model": payload.get("openaiModel", "gpt-4o-mini"),
                "apiKeyEnv": payload.get("openaiApiKeyEnv", "OPENAI_API_KEY"),
                "apiKey": "",
            },
        }
        CONFIG_DIR.mkdir(parents=True, exist_ok=True)
        MODEL_CONFIG_LOCAL_FILE.write_text(json.dumps(config, ensure_ascii=False, indent=2), encoding="utf-8")
        os.chmod(MODEL_CONFIG_LOCAL_FILE, 0o600)
        return {
            "message": "模型配置已保存，API Key 仅保留在当前服务进程内存中",
            "config": self.mask_model_config(self.config_for_display()),
            "activeProvider": self.active_writer_provider(),
            "engine": self.writer_engine_name(),
            "model": self.writer_model_name(),
        }

    def handle_model_config_test(self):
        text = self.call_configured_llm("请只输出：模型配置验证通过")
        if not text:
            raise ValueError("模型调用失败，请检查 API Key、请求地址和模型名称")
        return {
            "message": "模型配置验证通过",
            "engine": self.writer_engine_name(),
            "model": self.writer_model_name(),
            "outputPreview": text[:120],
        }

    def provider_priority_for(self, provider):
        mapping = {
            "anthropic": ["anthropic", "codex", "openai", "local"],
            "openai": ["openai", "anthropic", "codex", "local"],
            "codex": ["codex", "anthropic", "openai", "local"],
        }
        return mapping.get(provider, mapping["anthropic"])

    def mask_model_config(self, config):
        masked = json.loads(json.dumps(config, ensure_ascii=False))
        for section in [masked.get("env", {}), masked.get("anthropic", {}), masked.get("openai", {})]:
            for key in list(section.keys()):
                key_upper = key.upper()
                if key in {"apiKeyEnv"}:
                    continue
                if any(token in key_upper for token in ["APIKEY", "AUTH_TOKEN", "SECRET"]) and section.get(key):
                    section[key] = "********"
        return masked

    def parse_json_body(self):
        length = int(self.headers.get("Content-Length", "0") or 0)
        if length <= 0:
            return {}
        if length > MAX_JSON_BODY_BYTES:
            raise ValueError("JSON 请求体过大")
        return json.loads(self.rfile.read(length).decode("utf-8"))

    def config_for_display(self):
        config = load_model_config()
        if RUNTIME_SECRETS["anthropic"]:
            config.setdefault("env", {})["ANTHROPIC_AUTH_TOKEN"] = "********"
        if RUNTIME_SECRETS["openai"]:
            config.setdefault("openai", {})["apiKey"] = "********"
        return config

    def validate_base_url(self, raw_value, provider, allow_blank=False):
        value = (raw_value or "").strip()
        if not value:
            if allow_blank:
                return ""
            raise ValueError("模型基础地址不能为空")

        parsed = urlsplit(value)
        if parsed.scheme != "https":
            raise ValueError("模型基础地址必须使用 https")
        if not parsed.netloc or parsed.username or parsed.password:
            raise ValueError("模型基础地址格式不合法")
        if parsed.query or parsed.fragment:
            raise ValueError("模型基础地址不能包含查询串或片段")

        hostname = (parsed.hostname or "").lower()
        if not hostname:
            raise ValueError("模型基础地址缺少主机名")
        if self.host_is_private(hostname):
            raise ValueError("模型基础地址不能指向内网或保留地址")

        if provider == "openai":
            if hostname not in ALLOWED_OPENAI_HOSTS and not hostname.endswith(ALLOWED_OPENAI_SUFFIXES):
                raise ValueError("OpenAI 基础地址不在允许列表内")
        elif provider == "anthropic":
            if hostname not in ALLOWED_ANTHROPIC_HOSTS:
                raise ValueError("Anthropic 基础地址不在允许列表内")

        path = parsed.path.rstrip("/")
        return f"https://{parsed.netloc}{path}"

    def host_is_private(self, hostname):
        try:
            infos = socket.getaddrinfo(hostname, None, proto=socket.IPPROTO_TCP)
        except socket.gaierror:
            return False
        for info in infos:
            address = info[4][0]
            try:
                ip = ipaddress.ip_address(address)
            except ValueError:
                continue
            if ip.is_private or ip.is_loopback or ip.is_link_local or ip.is_multicast or ip.is_reserved or ip.is_unspecified:
                return True
        return False

    def no_redirect_open(self, request, timeout=60):
        class NoRedirectHandler(urllib.request.HTTPRedirectHandler):
            def redirect_request(self, req, fp, code, msg, headers, newurl):
                return None

        opener = urllib.request.build_opener(NoRedirectHandler)
        return opener.open(request, timeout=timeout)

    def handle_latest_proposal(self):
        session_id = self.require_session()
        job_dirs = [
            path for path in OUTPUT_DIR.iterdir()
            if path.is_dir() and (path / "manifest.json").exists() and self.job_belongs_to_session(path, session_id)
        ]
        if not job_dirs:
            raise ValueError("暂无历史任务")

        job_dir = max(job_dirs, key=lambda path: (path / "manifest.json").stat().st_mtime)
        job_id = job_dir.name
        manifest = json.loads((job_dir / "manifest.json").read_text(encoding="utf-8"))
        pipeline_result = {}
        pipeline_path = job_dir / "pipeline-result.json"
        if pipeline_path.exists():
            pipeline_result = json.loads(pipeline_path.read_text(encoding="utf-8"))

        return {
            "message": "已恢复最近任务",
            "jobId": job_id,
            **self.build_proposal_detail(job_dir, manifest, pipeline_result),
        }

    def handle_proposal_detail(self, job_id):
        job_dir = (OUTPUT_DIR / safe_job_id(job_id)).resolve()
        if not str(job_dir).startswith(str(OUTPUT_DIR.resolve())) or not job_dir.is_dir():
            raise ValueError("任务不存在")

        manifest_path = job_dir / "manifest.json"
        if not manifest_path.exists():
            raise ValueError("任务清单不存在")

        manifest = json.loads(manifest_path.read_text(encoding="utf-8"))
        self.authorize_job_access(job_dir, manifest)
        pipeline_result = {}
        pipeline_path = job_dir / "pipeline-result.json"
        if pipeline_path.exists():
            pipeline_result = json.loads(pipeline_path.read_text(encoding="utf-8"))

        return {
            "message": "已恢复任务",
            "jobId": job_dir.name,
            **self.build_proposal_detail(job_dir, manifest, pipeline_result),
        }

    def handle_outline_chat(self, job_id):
        job_dir = (OUTPUT_DIR / safe_job_id(job_id)).resolve()
        if not str(job_dir).startswith(str(OUTPUT_DIR.resolve())) or not job_dir.is_dir():
            raise ValueError("任务不存在")
        self.authorize_job_access(job_dir)

        payload = self.parse_json_body()
        message = self.clean_text(payload.get("message", ""))
        if not message:
            raise ValueError("请输入大纲调整要求")

        status = self.read_compose_status(job_dir)
        with ACTIVE_COMPOSE_LOCK:
            compose_is_active = job_id in ACTIVE_COMPOSE_TASKS
        if compose_is_active and status.get("state") in {"queued", "running"}:
            raise ValueError("当前正在正式编写，请等待完成后再调整大纲")

        pipeline_path = job_dir / "pipeline-result.json"
        if not pipeline_path.exists():
            raise ValueError("任务缺少大纲结果，请先生成大纲")

        pipeline_result = json.loads(pipeline_path.read_text(encoding="utf-8"))
        outline = pipeline_result.get("outline", {}).get("sections", [])
        adjusted_outline = (
            self.adjust_outline_by_hard_rules(outline, message)
            or self.adjust_outline_with_llm(outline, message)
            or self.adjust_outline_locally(outline, message)
        )
        if not adjusted_outline:
            raise ValueError("大纲调整失败，请换一种更明确的表达")
        adjusted_outline = self.diversify_outline_titles(adjusted_outline)

        pipeline_result.setdefault("outline", {})["sections"] = adjusted_outline
        pipeline_result.setdefault("outline", {})["lastAdjustment"] = {
            "message": message,
            "updatedAt": datetime.now(timezone.utc).isoformat(),
            "engine": self.writer_engine_name(),
            "model": self.writer_model_name(),
        }
        pipeline_result.setdefault("outline", {}).setdefault("chatHistory", []).append({
            "role": "user",
            "message": message,
            "createdAt": datetime.now(timezone.utc).isoformat(),
        })
        pipeline_result.setdefault("outline", {}).setdefault("chatHistory", []).append({
            "role": "assistant",
            "message": f"已调整为 {len(adjusted_outline)} 个一级章节，并刷新子标题层级。",
            "createdAt": datetime.now(timezone.utc).isoformat(),
        })
        pipeline_path.write_text(json.dumps(pipeline_result, ensure_ascii=False, indent=2), encoding="utf-8")

        return {
            "message": "大纲已根据对话要求调整",
            "jobId": job_id,
            "outline": adjusted_outline,
            "pipelineUrl": f"/output/{job_id}/pipeline-result.json",
            "engine": self.writer_engine_name(),
            "model": self.writer_model_name(),
        }

    def adjust_outline_by_hard_rules(self, outline, message):
        if not self.is_two_section_outline_request(message):
            return []

        scoring_items = []
        for section in outline:
            scoring_items.extend(section.get("relatedScoringItems") or [])
        scoring_items = scoring_items[:12]

        return [
            {
                "level": 1,
                "title": "技术方案",
                "writingGoal": "围绕技术路线、能力响应、实施交付和验收控制形成可评分、可落地的技术方案内容。",
                "score": None,
                "children": [
                    self.outline_child("项目整体理解与关键点分析", ["项目背景与建设目标", "关键需求识别", "技术难点与风险判断"]),
                    self.outline_child("总体技术架构与部署方案", ["总体架构设计", "系统部署拓扑", "关键组件协同机制"]),
                    self.outline_child("核心能力响应方案", ["功能能力匹配", "性能与可靠性设计", "安全防护能力说明"]),
                    self.outline_child("安全策略与配置方案", ["策略规划原则", "访问控制与防护策略", "日志审计与告警联动"]),
                    self.outline_child("实施计划与交付路径", ["阶段划分与里程碑", "实施步骤安排", "交付物清单"]),
                    self.outline_child("测试验收与风险控制", ["测试验证方法", "验收指标与证据", "风险识别与处置"]),
                ],
                "relatedScoringItems": scoring_items,
            },
            {
                "level": 1,
                "title": "服务方案",
                "writingGoal": "围绕服务组织、人员保障、响应速度、故障处理和持续运营形成可执行、可考核的服务方案内容。",
                "score": None,
                "children": [
                    self.outline_child("人员配备", ["项目角色分工", "专业能力配置", "人员投入与备份机制"]),
                    self.outline_child("响应速度", ["服务响应分级", "响应时限承诺", "升级通报机制"]),
                    self.outline_child("服务组织计划", ["服务组织架构", "服务流程安排", "沟通协调机制"]),
                    self.outline_child("故障处理措施", ["故障定位流程", "应急处置步骤", "复盘改进机制"]),
                    self.outline_child("日常维护与巡检方案", ["巡检频率与内容", "维护作业记录", "隐患跟踪闭环"]),
                    self.outline_child("服务质量保障方案", ["质量管理指标", "服务过程监督", "满意度与改进机制"]),
                    self.outline_child("培训与知识转移方案", ["培训对象与课程", "培训实施安排", "资料移交与答疑支持"]),
                ],
                "relatedScoringItems": scoring_items,
            },
        ]

    def is_two_section_outline_request(self, message):
        return (
            "技术方案" in message
            and "服务方案" in message
            and any(keyword in message for keyword in ["两个大标题", "2个大标题", "限定两个", "限定2个", "只有两个"])
        )

    def outline_child(self, title, children):
        return {
            "level": 2,
            "title": title,
            "children": [
                {
                    "level": 3,
                    "title": child,
                    "children": []
                }
                for child in children
            ],
        }

    def adjust_outline_with_llm(self, outline, message):
        prompt = f"""
你是一名网络安全售前方案大纲编辑智能体。请根据用户要求调整投标方案大纲。

用户要求：
{message}

当前大纲 JSON：
{json.dumps(outline, ensure_ascii=False, indent=2)}

请只输出 JSON，不要解释。格式必须为：
{{
  "sections": [
    {{
      "level": 1,
      "title": "章节名称",
      "writingGoal": "章节写作目标",
      "score": null,
      "children": [
        {{
          "level": 2,
          "title": "二级标题",
          "children": [
            {{
              "level": 3,
              "title": "三级标题",
              "children": [{{"level": 4, "title": "四级标题"}}]
            }}
          ]
        }}
      ],
      "relatedScoringItems": []
    }}
  ]
}}

调整要求：
1. 避免章节名称雷同，优先体现项目差异、技术路线、交付成果和验收口径。
2. 可以增删改章节、调整顺序、细化二级/三级/四级标题。
3. 每个一级章节最多 7 个二级标题；三级标题用于展开方法、流程、交付物、风险控制。
4. 保留 relatedScoringItems 字段，没有就使用空数组。
5. 不要输出 Markdown。
"""
        text = self.call_configured_llm(prompt)
        return self.parse_outline_response(text)

    def parse_outline_response(self, text):
        if not text:
            return []
        candidate = text.strip()
        fenced = re.search(r"```(?:json)?\s*(.*?)```", candidate, flags=re.S)
        if fenced:
            candidate = fenced.group(1).strip()
        start = candidate.find("{")
        end = candidate.rfind("}")
        if start >= 0 and end > start:
            candidate = candidate[start : end + 1]
        try:
            data = json.loads(candidate)
        except json.JSONDecodeError:
            return []
        sections = data.get("sections") if isinstance(data, dict) else data
        if not isinstance(sections, list):
            return []
        return [self.normalize_outline_section(section, 1) for section in sections if isinstance(section, dict)]

    def normalize_outline_section(self, section, level):
        title = self.clean_text(section.get("title") or f"未命名章节")
        normalized = {
            "level": level,
            "title": title,
            "writingGoal": self.clean_text(section.get("writingGoal") or f"围绕“{title}”形成可交付、可验收的方案内容。"),
            "score": section.get("score"),
            "children": [
                self.normalize_outline_section(child, min(level + 1, 4))
                for child in section.get("children", [])
                if isinstance(child, dict) and self.clean_text(child.get("title", ""))
            ][:7 if level == 1 else 5],
            "relatedScoringItems": section.get("relatedScoringItems") if isinstance(section.get("relatedScoringItems"), list) else [],
        }
        if level > 1:
            normalized.pop("writingGoal", None)
            normalized.pop("score", None)
            normalized.pop("relatedScoringItems", None)
        return normalized

    def adjust_outline_locally(self, outline, message):
        updated = json.loads(json.dumps(outline, ensure_ascii=False))
        if any(keyword in message for keyword in ["去掉", "删除", "移除"]) and any(keyword in message for keyword in ["第一章", "第1章", "首章"]):
            updated = updated[1:]
        if any(keyword in message for keyword in ["更细", "细化", "更多层级", "丰富"]):
            for section in updated:
                for child in section.get("children", []):
                    child.setdefault("children", [
                        {"level": 3, "title": "方法与路径", "children": [{"level": 4, "title": "执行动作"}, {"level": 4, "title": "输出材料"}]},
                        {"level": 3, "title": "验收与闭环", "children": [{"level": 4, "title": "检查口径"}, {"level": 4, "title": "客户确认"}]},
                    ])
        return updated

    def diversify_outline_titles(self, sections):
        seen_in_outline = {}
        for section in sections:
            children = section.get("children", [])
            self.diversify_sibling_titles(children, section.get("title", ""))
            for child in children:
                self.diversify_outline_subtree(child, section.get("title", ""), seen_in_outline)
        return sections

    def diversify_outline_subtree(self, node, parent_title, seen):
        title = self.clean_text(node.get("title", ""))
        level = int(node.get("level", 2) or 2)
        key = (level, title)
        seen[key] = seen.get(key, 0) + 1
        if level >= 2 and seen[key] > 1:
            node["title"] = self.unique_outline_title(parent_title, title, seen[key])

        children = node.get("children", [])
        if children:
            self.diversify_sibling_titles(children, node.get("title", title))
            for child in children:
                self.diversify_outline_subtree(child, node.get("title", title), seen)

    def diversify_sibling_titles(self, nodes, parent_title):
        seen = {}
        for node in nodes:
            title = self.clean_text(node.get("title", ""))
            seen[title] = seen.get(title, 0) + 1
            if seen[title] > 1:
                node["title"] = self.unique_outline_title(parent_title, title, seen[title])

    def unique_outline_title(self, parent_title, title, index):
        parent = re.sub(r"[一二三四五六七八九十0-9、.．\s]+$", "", self.clean_text(parent_title))
        parent = parent[:12] or "本项"
        replacements = {
            "现状与目标": "现状识别",
            "实施思路": "实施路径",
            "输出成果": "交付成果",
            "方法与路径": "执行路径",
            "验收与闭环": "验收闭环",
            "执行动作": "关键动作",
            "输出材料": "交付材料",
            "检查口径": "验收口径",
            "客户确认": "确认机制",
        }
        base = replacements.get(title, title or f"细化要点{index}")
        if base.startswith(parent):
            return base
        return f"{parent}{base}"

    def build_proposal_detail(self, job_dir, manifest, pipeline_result):
        job_id = job_dir.name
        return {
            "manifest": manifest,
            "outline": pipeline_result.get("outline", {}).get("sections", []),
            "composeStatus": self.compose_status_for_response(job_id, job_dir),
            "manifestUrl": f"/output/{job_id}/manifest.json",
            "draftUrl": f"/output/{job_id}/proposal-draft.md",
            "finalDraftUrl": f"/output/{job_id}/proposal-final-draft.md" if (job_dir / "proposal-final-draft.md").exists() else "",
            "finalDocxUrl": f"/output/{job_id}/proposal-final.docx" if (job_dir / "proposal-final.docx").exists() else "",
            "pipelineUrl": f"/output/{job_id}/pipeline-result.json",
        }

    def handle_proposal_request(self):
        form = self.parse_multipart_form()
        config = self.parse_config(form)
        self.validate_required_files(form)
        session_id = self.require_session()

        job_id = datetime.now(timezone.utc).strftime("%Y%m%d%H%M%S") + "-" + uuid.uuid4().hex[:8]
        job_dir = OUTPUT_DIR / job_id
        upload_dir = job_dir / "uploads"
        job_dir.mkdir(parents=True, exist_ok=False)
        upload_dir.mkdir(parents=True, exist_ok=True)

        files = {
            "technicalSpecification": self.save_upload("technicalSpecification", form["technicalSpecification"], upload_dir),
            "scoringCriteria": self.save_upload("scoringCriteria", form["scoringCriteria"], upload_dir),
        }

        template_upload = form.getfirst("templateFile")
        if "templateFile" in form and getattr(form["templateFile"], "filename", ""):
            files["proposalTemplate"] = self.save_upload("templateFile", form["templateFile"], upload_dir)
            template_source = "uploaded"
        else:
            files["proposalTemplate"] = self.copy_default_template(job_dir)
            template_source = "default"

        manifest = {
            "jobId": job_id,
            "createdAt": datetime.now(timezone.utc).isoformat(),
            "projectName": config.get("projectName") or "未填写",
            "outputType": config.get("outputType"),
            "templateSource": template_source,
            "files": files,
            "headingRules": config.get("postProcessing", {}).get("headingRules", {}),
            "pagePlan": config.get("postProcessing", {}).get("pagePlan", {}),
            "status": "draft_created",
        }

        pipeline_result = run_pipeline(job_dir, manifest, config, BASE_DIR)
        manifest["pipeline"] = {
            "materials": pipeline_result["materials"]["agent"],
            "scoring": pipeline_result["scoring"]["agent"],
            "template": pipeline_result["template"]["agent"],
            "outline": pipeline_result["outline"]["agent"],
            "export": pipeline_result["export"]["agent"],
        }
        manifest["status"] = "pipeline_completed"

        manifest_path = job_dir / "manifest.json"
        manifest_path.write_text(json.dumps(manifest, ensure_ascii=False, indent=2), encoding="utf-8")
        (job_dir / ACCESS_CONTROL_FILE).write_text(json.dumps({"ownerSessionId": session_id}, ensure_ascii=False), encoding="utf-8")

        return {
            "message": "方案生成任务已创建",
            "jobId": job_id,
            "templateSource": template_source,
            "manifestUrl": f"/output/{job_id}/manifest.json",
            "draftUrl": f"/output/{job_id}/proposal-draft.md",
            "pipelineUrl": f"/output/{job_id}/pipeline-result.json",
            "outline": pipeline_result["outline"]["sections"],
            "manifest": manifest,
        }

    def handle_compose_request(self, job_id):
        job_dir = (OUTPUT_DIR / safe_job_id(job_id)).resolve()
        if not str(job_dir).startswith(str(OUTPUT_DIR.resolve())) or not job_dir.is_dir():
            raise ValueError("任务不存在，请先生成并确认大纲")

        manifest_path = job_dir / "manifest.json"
        pipeline_path = job_dir / "pipeline-result.json"
        if not manifest_path.exists() or not pipeline_path.exists():
            raise ValueError("任务缺少大纲结果，请重新生成大纲")

        manifest = json.loads(manifest_path.read_text(encoding="utf-8"))
        self.authorize_job_access(job_dir, manifest)

        status = self.read_compose_status(job_dir)

        with ACTIVE_COMPOSE_LOCK:
            if job_id in ACTIVE_COMPOSE_TASKS:
                return status
            ACTIVE_COMPOSE_TASKS.add(job_id)

        initial_status = {
            "message": "已确认大纲，后台写作任务已启动",
            "jobId": job_id,
            "state": "queued",
            "stage": "queued",
            "progress": 0,
            "currentSection": "",
            "currentSubtitle": "",
            "startedAt": datetime.now(timezone.utc).isoformat(),
            "updatedAt": datetime.now(timezone.utc).isoformat(),
            "engine": self.writer_engine_name(),
            "model": self.writer_model_name(),
            "llmEnabled": self.llm_enabled(),
            "warning": self.writer_warning(),
            "manifestUrl": f"/output/{job_id}/manifest.json",
            "pipelineUrl": f"/output/{job_id}/pipeline-result.json",
        }
        self.write_compose_status(job_dir, initial_status)
        thread = threading.Thread(target=self.run_compose_task, args=(job_id,), daemon=True)
        thread.start()
        return initial_status

    def handle_compose_status(self, job_id):
        job_dir = (OUTPUT_DIR / safe_job_id(job_id)).resolve()
        if not str(job_dir).startswith(str(OUTPUT_DIR.resolve())) or not job_dir.is_dir():
            raise ValueError("任务不存在")
        self.authorize_job_access(job_dir)
        return self.compose_status_for_response(job_id, job_dir)

    def run_compose_task(self, job_id):
        job_dir = OUTPUT_DIR / safe_job_id(job_id)
        try:
            manifest_path = job_dir / "manifest.json"
            pipeline_path = job_dir / "pipeline-result.json"
            manifest = json.loads(manifest_path.read_text(encoding="utf-8"))
            pipeline_result = json.loads(pipeline_path.read_text(encoding="utf-8"))
            outline = pipeline_result.get("outline", {}).get("sections", [])
            page_allocations = self.allocate_section_pages(outline, manifest.get("pagePlan") or {})
            final_path = job_dir / "proposal-final-draft.md"
            final_docx_path = job_dir / "proposal-final.docx"

            self.update_compose_status(job_dir, {
                "state": "running",
                "stage": "planning",
                "progress": 3,
                "message": "正在规划章节篇幅与写作任务",
                "engine": self.writer_engine_name(),
                "model": self.writer_model_name(),
                "pageAllocations": page_allocations,
            })
            final_path.write_text(self.build_final_draft(manifest, outline, page_allocations), encoding="utf-8")

            def progress_callback(payload):
                self.update_compose_status(job_dir, payload)

            page_control = self.build_final_docx(
                manifest,
                pipeline_result,
                outline,
                page_allocations,
                final_docx_path,
                progress_callback=progress_callback,
            )

            manifest["status"] = "composition_completed"
            manifest["finalDraft"] = str(final_path.relative_to(BASE_DIR))
            manifest["finalDocx"] = str(final_docx_path.relative_to(BASE_DIR))
            manifest["pageAllocations"] = page_allocations
            manifest["pageControl"] = page_control
            manifest_path.write_text(json.dumps(manifest, ensure_ascii=False, indent=2), encoding="utf-8")

            self.update_compose_status(job_dir, {
                "state": "completed",
                "stage": "done",
                "progress": 100,
                "message": "正式方案编写完成，可下载 Word 方案",
                "manifestUrl": f"/output/{job_id}/manifest.json",
                "draftUrl": f"/output/{job_id}/proposal-draft.md",
                "finalDraftUrl": f"/output/{job_id}/proposal-final-draft.md",
                "finalDocxUrl": f"/output/{job_id}/proposal-final.docx",
                "pipelineUrl": f"/output/{job_id}/pipeline-result.json",
                "pageAllocations": page_allocations,
                "pageControl": page_control,
                "completedAt": datetime.now(timezone.utc).isoformat(),
            })
        except Exception as exc:
            logging.exception("Compose task failed for %s", job_id)
            self.update_compose_status(job_dir, {
                "state": "failed",
                "stage": "failed",
                "message": "正式编写失败，请检查输入材料或查看服务器日志",
            })
        finally:
            with ACTIVE_COMPOSE_LOCK:
                ACTIVE_COMPOSE_TASKS.discard(job_id)

    def read_compose_status(self, job_dir):
        status_path = job_dir / COMPOSE_STATUS_FILE
        if status_path.exists():
            return json.loads(status_path.read_text(encoding="utf-8"))
        return {
            "state": "idle",
            "stage": "idle",
            "progress": 0,
            "message": "等待确认大纲并开始编写",
            "engine": self.writer_engine_name(),
            "model": self.writer_model_name(),
            "llmEnabled": self.llm_enabled(),
        }

    def compose_status_for_response(self, job_id, job_dir):
        status = self.read_compose_status(job_dir)
        status.pop("error", None)
        status.setdefault("engine", self.writer_engine_name())
        status.setdefault("model", self.writer_model_name())
        if status.get("state") in {"queued", "running"}:
            with ACTIVE_COMPOSE_LOCK:
                is_active = job_id in ACTIVE_COMPOSE_TASKS
            if not is_active:
                status = {
                    **status,
                    "state": "failed",
                    "stage": "interrupted",
                    "message": "后台写作任务已中断，请重新点击确认大纲并开始编写",
                    "progress": status.get("progress", 0),
                }
        return status

    def write_compose_status(self, job_dir, status):
        status["updatedAt"] = datetime.now(timezone.utc).isoformat()
        (job_dir / COMPOSE_STATUS_FILE).write_text(json.dumps(status, ensure_ascii=False, indent=2), encoding="utf-8")

    def update_compose_status(self, job_dir, patch):
        status = self.read_compose_status(job_dir)
        status.update(patch)
        self.write_compose_status(job_dir, status)

    def allocate_section_pages(self, outline, page_plan):
        total_pages = page_plan.get("totalPages") or max(len(outline) * 3, 12)
        manual_pages = {
            item.get("chapterTitle"): item.get("pages")
            for item in page_plan.get("chapterPages", [])
            if item.get("chapterTitle") and item.get("pages")
        }
        allocations = []
        remaining_sections = []
        assigned = 0

        for section in outline:
            title = section.get("title")
            manual = manual_pages.get(title)
            if manual:
                pages = max(1, int(manual))
                assigned += pages
                allocations.append({"title": title, "pages": pages, "source": "manual"})
            else:
                remaining_sections.append(section)

        remaining_pages = max(len(remaining_sections), int(total_pages) - assigned)
        weights = [max(1, int(section.get("score") or 3)) for section in remaining_sections]
        weight_total = sum(weights) or 1
        auto_pages = []

        for index, section in enumerate(remaining_sections):
            if index == len(remaining_sections) - 1:
                pages = remaining_pages - sum(auto_pages)
            else:
                pages = max(1, round(remaining_pages * weights[index] / weight_total))
            auto_pages.append(max(1, pages))
            allocations.append({"title": section.get("title"), "pages": max(1, pages), "source": "auto"})

        order = {section.get("title"): index for index, section in enumerate(outline)}
        allocations.sort(key=lambda item: order.get(item["title"], 999))
        return allocations

    def build_final_draft(self, manifest, outline, page_allocations):
        page_map = {item["title"]: item["pages"] for item in page_allocations}
        lines = [
            f"# {manifest['projectName']} 正式方案初稿",
            "",
            "## 编写说明",
            "",
            "本文件基于用户确认的大纲生成，并按照篇幅规划分配章节内容。Word 正式稿请下载 proposal-final.docx。",
            "",
        ]

        for section in outline:
            lines.extend(
                [
                    f"## {section.get('title', '未命名章节')}",
                    "",
                    f"目标页数：{page_map.get(section.get('title'), 1)} 页",
                    "",
                    section.get("writingGoal", "围绕本章节生成可交付内容。"),
                    "",
                    "### 编写要点",
                    "",
                ]
            )
            scoring_items = section.get("relatedScoringItems") or []
            if scoring_items:
                lines.append("- 需要重点响应以下评分项：")
                for item in scoring_items:
                    lines.append(f"  - {item}")
            else:
                lines.append("- 根据技术规范书和默认模板补充章节内容。")

            lines.extend(
                [
                    "- 保持与模板标题层级、段落风格和篇幅规划一致。",
                    "",
                    "### 正文草稿",
                    "",
                    *self.compose_section_markdown(section, page_map.get(section.get("title"), 1), use_writer=False),
                    "",
                ]
            )

        return "\n".join(lines)

    def build_final_docx(self, manifest, pipeline_result, outline, page_allocations, target_path, progress_callback=None):
        template_path = BASE_DIR / manifest["files"]["proposalTemplate"]["path"]
        document = Document(template_path)
        self.clear_document(document)
        page_control = {
            "targetCharsPerPage": TARGET_CHARS_PER_PAGE,
            "sections": [],
        }

        page_map = {item["title"]: item["pages"] for item in page_allocations}
        total_target_pages = sum(item.get("pages", 0) for item in page_allocations)
        density_factor = self.content_density_factor(total_target_pages)
        page_control["densityFactor"] = density_factor
        technical_signals = pipeline_result.get("materials", {}).get("technicalSpecification", {}).get("requirementSignals", [])
        total_children = sum(len(section.get("children") or [{"title": "章节内容"}]) for section in outline) or 1
        completed_children = 0

        for section_index, section in enumerate(outline, start=1):
            title = section.get("title", "未命名章节")
            target_pages = page_map.get(title, 1)
            if progress_callback:
                progress_callback({
                    "state": "running",
                    "stage": "section",
                    "message": f"正在编写第 {section_index}/{len(outline)} 章：{title}",
                    "currentSection": title,
                    "currentSubtitle": "",
                    "progress": max(5, round(completed_children * 90 / total_children)),
                })
            section_start_chars = self.document_text_length(document)
            document.add_heading(title, level=1)
            self.add_section_table(document, title)

            children = section.get("children") or [{"title": "章节内容"}]
            content_pages = self.section_content_page_budget(title, target_pages)
            child_budgets = self.allocate_child_char_budgets(children, content_pages, density_factor)
            child_stats = []
            for child, child_budget in zip(children, child_budgets):
                subtitle = child.get("title", "章节内容")
                if progress_callback:
                    progress_callback({
                        "state": "running",
                        "stage": "writing",
                        "message": f"正在编写：{title} / {subtitle}",
                        "currentSection": title,
                        "currentSubtitle": subtitle,
                        "engine": self.writer_engine_name(),
                        "model": self.writer_model_name(),
                        "progress": max(6, round(completed_children * 90 / total_children)),
                    })
                child_start_chars = self.document_text_length(document)
                document.add_heading(subtitle, level=2)
                child_progress = max(6, round(completed_children * 90 / total_children))
                for paragraph in self.compose_paragraphs(
                    title,
                    subtitle,
                    section,
                    technical_signals,
                    target_chars=child_budget,
                    progress_callback=progress_callback,
                    base_progress=child_progress,
                ):
                    self.add_body_paragraph(document, paragraph)
                self.add_outline_substructure(document, child, title, section, technical_signals)
                child_chars = self.document_text_length(document) - child_start_chars
                completed_children += 1
                child_stats.append({
                    "title": subtitle,
                    "targetChars": child_budget,
                    "actualChars": child_chars,
                    "estimatedPages": round(child_chars / TARGET_CHARS_PER_PAGE, 2),
                })

            section_chars = self.document_text_length(document) - section_start_chars
            page_control["sections"].append({
                "title": title,
                "targetPages": target_pages,
                "contentBudgetPages": content_pages,
                "targetChars": round(content_pages * TARGET_CHARS_PER_PAGE * density_factor),
                "actualChars": section_chars,
                "estimatedPages": round(section_chars / TARGET_CHARS_PER_PAGE, 2),
                "children": child_stats,
            })

        if progress_callback:
            progress_callback({
                "state": "running",
                "stage": "export",
                "message": "正在套用模板并生成 Word 正式稿",
                "currentSection": "",
                "currentSubtitle": "",
                "progress": 95,
            })
        document.save(target_path)
        total_chars = self.document_text_length(document)
        page_control["totalActualChars"] = total_chars
        page_control["totalEstimatedPages"] = round(total_chars / TARGET_CHARS_PER_PAGE, 2)
        page_control["totalTargetPages"] = sum(item.get("pages", 0) for item in page_allocations)
        return page_control

    def add_body_paragraph(self, document, text=""):
        paragraph = document.add_paragraph(text)
        paragraph.paragraph_format.first_line_indent = BODY_FIRST_LINE_INDENT
        return paragraph

    def add_outline_substructure(self, document, node, section_title, section=None, technical_signals=None):
        children = node.get("children") or []
        if not children:
            return
        section = section or {}
        technical_signals = technical_signals or []
        scoring_items = section.get("relatedScoringItems") or []
        scoring_text = self.clean_text(scoring_items[0]) if scoring_items else "围绕评分标准要求进行针对性响应。"
        context = self.context_summary(technical_signals)
        for child in children:
            level = min(max(int(child.get("level", 3)), 3), 4)
            child_title = child.get("title", "细化内容")
            document.add_heading(child_title, level=level)
            leaf_titles = [leaf.get("title", "") for leaf in child.get("children", []) if leaf.get("title")]
            if leaf_titles:
                for paragraph in self.compose_subsection_paragraphs(section_title, node.get("title", ""), child_title, scoring_text, context):
                    self.add_body_paragraph(document, paragraph)
                self.add_bullet_list(document, leaf_titles)
                self.add_check_table(document, child.get("title", ""), leaf_titles[:3])
            else:
                for paragraph in self.compose_subsection_paragraphs(section_title, node.get("title", ""), child_title, scoring_text, context):
                    self.add_body_paragraph(document, paragraph)
        self.add_callout(document, f"{node.get('title', '本小节')}应与评分条款、交付物清单和验收记录形成对应关系，避免只描述原则而缺少过程支撑。")

    def compose_subsection_paragraphs(self, section_title, parent_title, child_title, scoring_text, context):
        profile = self.subtitle_profile(section_title, child_title)
        context_points = self.context_points(context)
        context_point = self.choose_context_point(section_title, child_title, context_points, 0) if context_points else profile["problem"]
        scoring_brief = self.scoring_brief(scoring_text)
        return [
            (
                f"在“{parent_title}”下设置“{child_title}”，主要用于把{profile['object']}落到可执行层面。"
                f"本部分围绕{profile['action']}展开，结合评分关注的“{scoring_brief}”，明确服务方需要完成的动作、形成的材料以及客户侧确认方式。"
            ),
            (
                f"执行过程中，项目组先围绕“{context_point}”确认输入条件和工作边界，再安排责任角色推进具体任务。"
                f"每项动作均通过{profile['evidence']}留痕，避免三级小节只保留标题而缺少可复核内容。"
            ),
            (
                f"验收方面，“{child_title}”以{profile['metric']}作为检查口径。"
                f"若发现{profile['risk']}等风险，项目组应及时记录影响范围、处置措施和复核结论，并在阶段汇报中向客户确认闭环状态。"
            ),
        ]

    def add_bullet_list(self, document, items):
        style = self.first_existing_style(document, ["List Bullet", "List Paragraph", "列表段落"])
        for item in items[:5]:
            if style:
                paragraph = document.add_paragraph(style=style)
                paragraph.add_run(item)
            else:
                paragraph = document.add_paragraph(f"• {item}")
                paragraph.paragraph_format.first_line_indent = None
                paragraph.paragraph_format.left_indent = Pt(18)

    def first_existing_style(self, document, names):
        for name in names:
            try:
                document.styles[name]
                return name
            except KeyError:
                continue
        return ""

    def add_callout(self, document, text):
        table = document.add_table(rows=1, cols=1)
        table.style = "Table Grid"
        cell = table.rows[0].cells[0]
        cell.text = f"关键控制点：{text}"
        for paragraph in cell.paragraphs:
            paragraph.paragraph_format.first_line_indent = None
        document.add_paragraph("")

    def add_check_table(self, document, title, items):
        table = document.add_table(rows=1, cols=3)
        table.style = "Table Grid"
        headers = ["检查项", "输出材料", "确认口径"]
        for index, header in enumerate(headers):
            table.rows[0].cells[index].text = header
        for item in items:
            cells = table.add_row().cells
            cells[0].text = item
            cells[1].text = f"{title}记录"
            cells[2].text = "客户确认、过程留痕、结果可复核"
        document.add_paragraph("")

    def clear_document(self, document):
        body = document._element.body
        for child in list(body):
            if child.tag.endswith("sectPr"):
                continue
            body.remove(child)

    def allocate_child_pages(self, children, section_pages):
        count = max(1, len(children))
        base = max(1, section_pages // count)
        allocations = [base for _ in children]
        for index in range(max(0, section_pages - sum(allocations))):
            allocations[index % count] += 1
        return allocations

    def allocate_child_char_budgets(self, children, section_content_pages, density_factor=1.0):
        count = max(1, len(children))
        total_chars = max(MIN_CHILD_CHARS * count, round(section_content_pages * TARGET_CHARS_PER_PAGE * density_factor))
        base = total_chars // count
        budgets = [max(MIN_CHILD_CHARS, base) for _ in children]
        remainder = total_chars - sum(budgets)
        index = 0
        while remainder > 0 and budgets:
            add = min(80, remainder)
            budgets[index % count] += add
            remainder -= add
            index += 1
        return budgets

    def section_content_page_budget(self, section_title, section_pages):
        overhead = 1.0 if self.has_section_table(section_title) else 0.6
        overhead = min(overhead, max(0.0, section_pages - 0.35))
        return max(0.35, section_pages - overhead)

    def content_density_factor(self, total_target_pages):
        if total_target_pages <= 80:
            return 0.78
        if total_target_pages <= 150:
            return 0.94
        if total_target_pages <= 250:
            return 0.97
        return 1.0

    def compose_section_markdown(self, section, pages, use_writer=True):
        lines = []
        children = section.get("children") or [{"title": "章节内容"}]
        for child in children:
            lines.append(f"#### {child.get('title')}")
            lines.append("")
            if use_writer:
                child_budget = max(MIN_CHILD_CHARS, round((pages * TARGET_CHARS_PER_PAGE) / max(1, len(children))))
                lines.extend(self.compose_paragraphs(section.get("title", ""), child.get("title", ""), section, [], target_chars=child_budget))
            else:
                lines.extend([
                    "- 正式正文将在确认编写后由后台写作引擎逐小节生成。",
                    "- 本小节需结合评分项、技术规范书和模板格式进行展开。",
                    "- 生成过程中可在页面查看当前章节、小节和写作进度。",
                ])
            lines.append("")
        return lines

    def compose_paragraphs(self, section_title, subtitle, section, technical_signals, target_chars=None, progress_callback=None, base_progress=0):
        scoring_items = section.get("relatedScoringItems") or []
        scoring_text = self.clean_text(scoring_items[0]) if scoring_items else "围绕评分标准要求进行针对性响应。"
        context = self.context_summary(technical_signals)
        target_chars = max(MIN_CHILD_CHARS, int(target_chars or TARGET_CHARS_PER_PAGE))
        paragraphs = self.compose_paragraphs_with_llm(
            section_title,
            subtitle,
            scoring_text,
            context,
            target_chars,
            progress_callback=progress_callback,
            base_progress=base_progress,
        )
        if paragraphs and self.paragraphs_text_length(paragraphs) >= target_chars * 0.9:
            return paragraphs

        base_paragraphs = self.section_paragraph_bank(section_title, subtitle, scoring_text, context)
        for paragraph in base_paragraphs:
            self.append_unique_paragraph(paragraphs, paragraph)
            if len(paragraphs) >= 7 and self.paragraphs_text_length(paragraphs) >= target_chars * 0.88:
                return paragraphs

        for paragraph in self.expanded_writing_paragraphs(section_title, subtitle, scoring_text, context):
            self.append_unique_paragraph(paragraphs, paragraph)
            if len(paragraphs) >= 7 and self.paragraphs_text_length(paragraphs) >= target_chars:
                break

        index = 0
        while self.paragraphs_text_length(paragraphs) < target_chars:
            before = len(paragraphs)
            self.append_unique_paragraph(
                paragraphs,
                self.long_form_paragraph(section_title, subtitle, scoring_text, context, index),
            )
            index += 1
            if len(paragraphs) == before and index > 80:
                break

        return paragraphs

    def compose_paragraphs_with_llm(self, section_title, subtitle, scoring_text, context, target_chars, progress_callback=None, base_progress=0):
        if not self.llm_enabled():
            return []

        profile = self.subtitle_profile(section_title, subtitle)
        context_points = self.context_points(context)
        relevant_points = [
            self.choose_context_point(section_title, subtitle, context_points, index)
            for index in range(min(4, len(context_points)))
        ] if context_points else []
        target_chars = min(max(target_chars, 900), 2200)
        prompt = f"""
你是一名资深网络安全售前方案工程师，正在编写投标技术方案。

项目章节：{section_title}
当前小节：{subtitle}
目标字数：约 {target_chars} 个中文字符

评分项摘要：
{self.scoring_brief(scoring_text)}

技术规范书相关材料：
{chr(10).join('- ' + point for point in relevant_points) if relevant_points else '- 暂未匹配到细化条款，请围绕评分标准写作并预留可核对口径。'}

写作对象：
- 服务对象：{profile['object']}
- 执行主线：{profile['action']}
- 证明材料：{profile['evidence']}
- 风险关注：{profile['risk']}
- 验收口径：{profile['metric']}

请直接输出正式方案正文，要求：
1. 使用中文投标方案文风，不要解释你在做什么。
2. 输出 5 到 8 个自然段，每段 70 到 150 字，句式要有长短变化。
3. 必须结合评分项和技术规范书，至少写清一个执行动作、一个输出材料和一个验收口径。
4. 不要输出标题、编号、Markdown、项目符号。
5. 避免“高度重视、全面保障、持续优化”等空泛堆词，优先写项目场景、责任角色、交付记录和客户确认方式。
6. 段落之间用换行分隔。
"""
        text = self.call_configured_llm(
            prompt,
            progress_callback=progress_callback,
            section_title=section_title,
            subtitle=subtitle,
            base_progress=base_progress,
        )
        paragraphs = []
        for line in text.splitlines():
            line = self.clean_text(line)
            if line and not line.startswith(("#", "-", "1.", "1、")):
                self.append_unique_paragraph(paragraphs, line)
        return paragraphs

    def call_configured_llm(self, prompt, progress_callback=None, section_title="", subtitle="", base_progress=0):
        for provider in self.writer_provider_priority():
            if provider == "anthropic" and self.anthropic_enabled():
                text = self.call_anthropic_messages(prompt)
                if text:
                    return text
            if provider == "codex" and self.codex_enabled():
                text = self.call_codex_exec(
                    prompt,
                    progress_callback=progress_callback,
                    section_title=section_title,
                    subtitle=subtitle,
                    base_progress=base_progress,
                )
                if text:
                    return text
            if provider == "openai" and self.openai_enabled():
                text = self.call_openai_chat(prompt)
                if text:
                    return text
        return ""

    def call_codex_exec(self, prompt, progress_callback=None, section_title="", subtitle="", base_progress=0):
        if not self.codex_enabled():
            return ""
        output_path = OUTPUT_DIR / f"codex-last-message-{uuid.uuid4().hex}.txt"
        codex_model = self.codex_model()
        command = [
            "codex",
            "exec",
            "--sandbox",
            "read-only",
            "--skip-git-repo-check",
            "--output-last-message",
            str(output_path),
        ]
        if codex_model:
            command.extend(["--model", codex_model])
        command.append("-")
        try:
            started = time.monotonic()
            process = subprocess.Popen(
                command,
                text=True,
                cwd=str(BASE_DIR),
                env=self.model_subprocess_env(),
                stdin=subprocess.PIPE,
                stdout=subprocess.DEVNULL,
                stderr=subprocess.DEVNULL,
            )
            if process.stdin:
                process.stdin.write(prompt)
                process.stdin.close()

            while process.poll() is None:
                elapsed = round(time.monotonic() - started)
                if elapsed > self.codex_timeout_seconds():
                    process.kill()
                    process.wait(timeout=5)
                    return ""
                if progress_callback:
                    progress_callback({
                        "state": "running",
                        "stage": "model",
                        "message": f"正在调用模型生成正文：{section_title} / {subtitle}（已耗时 {elapsed} 秒）",
                        "currentSection": section_title,
                        "currentSubtitle": subtitle,
                        "progress": min(94, max(base_progress, base_progress + min(3, elapsed // 30))),
                        "engine": self.writer_engine_name(),
                        "model": self.writer_model_name(),
                        "modelElapsedSeconds": elapsed,
                        "heartbeatAt": datetime.now(timezone.utc).isoformat(),
                    })
                time.sleep(5)

            if output_path.exists():
                return output_path.read_text(encoding="utf-8", errors="ignore").strip()
        except (subprocess.SubprocessError, OSError):
            return ""
        finally:
            try:
                output_path.unlink(missing_ok=True)
            except OSError:
                pass
        return ""

    def call_anthropic_messages(self, prompt):
        config = self.anthropic_config()
        api_key = self.anthropic_api_key()
        base_url = self.anthropic_base_url()
        model = self.anthropic_model()
        if not (api_key and base_url and model):
            return ""
        payload = {
            "model": model,
            "max_tokens": int(config.get("maxTokens", 1800)),
            "messages": [
                {"role": "user", "content": prompt},
            ],
        }
        request = urllib.request.Request(
            f"{base_url}/messages",
            data=json.dumps(payload, ensure_ascii=False).encode("utf-8"),
            headers={
                "Authorization": f"Bearer {api_key}",
                "x-api-key": api_key,
                "anthropic-version": str(config.get("version", "2023-06-01")),
                "Content-Type": "application/json",
            },
            method="POST",
        )
        try:
            with self.no_redirect_open(request, timeout=int(config.get("timeoutSeconds", 90))) as response:
                data = json.loads(response.read().decode("utf-8"))
            content = data.get("content", [])
            return "\n".join(item.get("text", "") for item in content if isinstance(item, dict)).strip()
        except (urllib.error.URLError, KeyError, IndexError, json.JSONDecodeError, TimeoutError):
            return ""

    def call_openai_chat(self, prompt):
        api_key = self.openai_api_key()
        if not api_key:
            return ""
        config = self.openai_config()
        payload = {
            "model": self.openai_model(),
            "messages": [
                {"role": "system", "content": "你是严谨、专业、懂网络安全售前投标的中文方案写作智能体。"},
                {"role": "user", "content": prompt},
            ],
            "temperature": 0.35,
        }
        request = urllib.request.Request(
            f"{self.openai_base_url()}/chat/completions",
            data=json.dumps(payload, ensure_ascii=False).encode("utf-8"),
            headers={
                "Authorization": f"Bearer {api_key}",
                "Content-Type": "application/json",
            },
            method="POST",
        )
        try:
            with self.no_redirect_open(request, timeout=int(config.get("timeoutSeconds", 90))) as response:
                data = json.loads(response.read().decode("utf-8"))
            return data["choices"][0]["message"]["content"]
        except (urllib.error.URLError, KeyError, IndexError, json.JSONDecodeError):
            return ""

    def codex_enabled(self):
        return env_bool("CODEX_WRITER_ENABLED", bool(self.codex_config().get("enabled", True))) and shutil.which("codex") is not None

    def anthropic_enabled(self):
        return bool(self.anthropic_config().get("enabled", True)) and bool(self.anthropic_api_key())

    def openai_enabled(self):
        return bool(self.openai_config().get("enabled", True)) and bool(self.openai_api_key())

    def current_model_config(self):
        return load_model_config()

    def current_config_env(self):
        env = self.current_model_config().get("env", {})
        return {str(key): str(value) for key, value in env.items()}

    def model_subprocess_env(self):
        env = os.environ.copy()
        env.update(self.current_config_env())
        return env

    def writer_config(self):
        return self.current_model_config().get("writer", {})

    def codex_config(self):
        return self.current_model_config().get("codex", {})

    def anthropic_config(self):
        return self.current_model_config().get("anthropic", {})

    def openai_config(self):
        return self.current_model_config().get("openai", {})

    def codex_model(self):
        return os.environ.get("CODEX_MODEL") or self.current_config_env().get("CODEX_MODEL") or self.codex_config().get("model", "")

    def codex_timeout_seconds(self):
        return int(os.environ.get("CODEX_TIMEOUT_SECONDS") or self.current_config_env().get("CODEX_TIMEOUT_SECONDS") or self.codex_config().get("timeoutSeconds", 240))

    def anthropic_model(self):
        env = self.current_config_env()
        return os.environ.get("ANTHROPIC_MODEL") or env.get("ANTHROPIC_MODEL") or self.anthropic_config().get("model", "")

    def anthropic_api_key(self):
        env = self.current_config_env()
        key_env = self.anthropic_config().get("apiKeyEnv", "ANTHROPIC_AUTH_TOKEN")
        return os.environ.get(key_env) or RUNTIME_SECRETS["anthropic"] or env.get(key_env) or self.anthropic_config().get("apiKey", "")

    def anthropic_base_url(self):
        value = os.environ.get("ANTHROPIC_BASE_URL") or self.current_config_env().get("ANTHROPIC_BASE_URL") or self.anthropic_config().get("baseUrl", "")
        return self.validate_base_url(value, "anthropic", allow_blank=True)

    def openai_model(self):
        return os.environ.get("OPENAI_MODEL") or self.current_config_env().get("OPENAI_MODEL") or self.openai_config().get("model", "gpt-4o-mini")

    def openai_base_url(self):
        value = os.environ.get("OPENAI_BASE_URL") or self.current_config_env().get("OPENAI_BASE_URL") or self.openai_config().get("baseUrl", "https://api.openai.com/v1")
        return self.validate_base_url(value, "openai")

    def openai_api_key(self):
        env = self.current_config_env()
        key_env = os.environ.get("OPENAI_API_KEY_ENV") or env.get("OPENAI_API_KEY_ENV") or self.openai_config().get("apiKeyEnv", "OPENAI_API_KEY")
        return os.environ.get(key_env) or RUNTIME_SECRETS["openai"] or env.get(key_env) or self.openai_config().get("apiKey", "") or os.environ.get("OPENAI_API_KEY", "")

    def llm_enabled(self):
        return self.anthropic_enabled() or self.codex_enabled() or self.openai_enabled()

    def writer_provider_priority(self):
        priority = self.writer_config().get("providerPriority", ["codex", "openai", "local"])
        return [str(item).strip().lower() for item in priority if str(item).strip()]

    def active_writer_provider(self):
        for provider in self.writer_provider_priority():
            if provider == "anthropic" and self.anthropic_enabled():
                return "anthropic"
            if provider == "codex" and self.codex_enabled():
                return "codex"
            if provider == "openai" and self.openai_enabled():
                return "openai"
            if provider == "local":
                return "local"
        return "local"

    def writer_engine_name(self):
        provider = self.active_writer_provider()
        if provider == "anthropic":
            return "Anthropic-compatible API"
        if provider == "codex":
            return "Codex CLI"
        if provider == "openai":
            return "OpenAI API"
        return "local-fallback-writer"

    def writer_model_name(self):
        provider = self.active_writer_provider()
        if provider == "anthropic":
            return self.anthropic_model()
        if provider == "codex":
            codex_model = self.codex_model()
            if codex_model:
                return f"{codex_model}（由配置或 CODEX_MODEL 指定）"
            return "Codex CLI 当前默认模型（跟随本机 Codex 登录/配置）"
        if provider == "openai":
            return self.openai_model()
        return "local-template-writer"

    def writer_warning(self):
        if self.active_writer_provider() != "local":
            return ""
        return "未检测到可用模型配置，当前为本地降级写作模式。"

    def paragraphs_text_length(self, paragraphs):
        return sum(len(paragraph) for paragraph in paragraphs)

    def append_unique_paragraph(self, paragraphs, paragraph):
        paragraph = self.clean_text(paragraph)
        if not paragraph:
            return
        candidate = self.normalized_paragraph(paragraph)
        for existing in paragraphs:
            if candidate == self.normalized_paragraph(existing):
                return
            if self.paragraph_similarity(candidate, self.normalized_paragraph(existing)) > 0.78:
                return
        paragraphs.append(paragraph)

    def normalized_paragraph(self, paragraph):
        return re.sub(r"[，。；：、“”《》（）()\\s]+", "", paragraph)

    def paragraph_similarity(self, left, right):
        if not left or not right:
            return 0
        left_tokens = {left[index : index + 2] for index in range(max(1, len(left) - 1))}
        right_tokens = {right[index : index + 2] for index in range(max(1, len(right) - 1))}
        return len(left_tokens & right_tokens) / max(1, len(left_tokens | right_tokens))

    def expanded_writing_paragraphs(self, section_title, subtitle, scoring_text, context):
        profile = self.subtitle_profile(section_title, subtitle)
        scoring_brief = self.scoring_brief(scoring_text)
        context_points = self.context_points(context)
        paragraphs = [
            f"在需求承接上，“{subtitle}”不作为孤立条目处理，而是承接“{section_title}”的评审重点，围绕{profile['object']}展开，以{profile['action']}为执行主线，并以{profile['evidence']}作为证明材料。",
            f"在场景理解上，“{subtitle}”首先明确客户侧需要解决的实际问题：{profile['problem']}。方案正文将围绕该问题说明现状约束、服务介入方式、过程控制点和结果确认方式，使评审人员能够看到本小节与项目需求之间的对应关系。",
            f"在服务动作设计上，项目组将围绕“{subtitle}”以{profile['action']}为主线，拆分为准备、执行、复核、反馈和归档五个连续动作。每个动作均明确责任角色和输出材料，避免只给出原则性承诺。",
            f"在客户协同上，“{subtitle}”需要与客户接口人形成稳定互动。项目组将在事项发起、需求确认、执行反馈和结果验收环节保留沟通记录，确保客户能够及时掌握本小节服务进展并对关键结论进行确认。",
            f"在材料支撑上，“{subtitle}”对应的证明材料包括{profile['evidence']}。这些材料不是附属清单，而是正文论证的一部分，用于证明服务动作已经执行、过程可追溯、结果可复核。",
            f"在验收表达上，方案将把{profile['metric']}作为“{subtitle}”的检查依据，说明客户如何判断本小节工作是否完成、是否有效、是否满足评分要求。",
        ]

        if context_points:
            for index, point in enumerate(context_points[:4], start=1):
                paragraphs.append(
                    f"结合技术规范书第 {index} 类信息，“{subtitle}”将“{point}”转化为具体方案要求：一是明确该要求影响的服务环节，二是说明项目组采取的执行动作，三是给出可提交给客户确认的记录或报告。"
                )

        stages = [
            ("启动阶段", "完成需求口径统一、接口人确认和材料清单建立"),
            ("实施阶段", "按照计划开展服务动作，并在关键节点同步执行结果"),
            ("运行阶段", "持续跟踪问题、指标和客户反馈，及时调整服务安排"),
            ("总结阶段", "整理过程记录、复盘问题原因并形成改进建议"),
        ]
        for stage, action in stages:
            paragraphs.append(
                f"在{stage}，项目组围绕“{subtitle}”重点{action}。该阶段的工作不只关注任务是否完成，还关注完成质量、客户确认情况以及后续是否需要补充优化。"
            )

        dimensions = [
            ("边界控制", f"明确{profile['object']}的适用范围、触发条件和不适用场景，减少执行中的理解偏差。"),
            ("责任分工", f"将{profile['action']}拆解到项目负责人、技术执行人员、质量复核人员和客户接口人，保证每项工作有人推进、有人复核、有人确认。"),
            ("记录留痕", f"围绕{profile['evidence']}建立统一记录口径，记录时间、事项、处理人、结论和客户反馈。"),
            ("风险预防", f"针对{profile['risk']}设置预警条件和处置路径，避免风险扩大后才被动响应。"),
            ("质量复核", f"对{profile['metric']}进行阶段性检查，发现偏差时及时修正服务动作和材料表达。"),
            ("成果反馈", f"将执行进展、发现问题、处理建议和后续计划整理为客户可阅读的反馈内容。"),
            ("验收映射", f"把评分条款中的关键词映射到正文段落、过程记录和交付物，便于评审时快速定位响应内容。"),
            ("持续优化", f"根据服务记录和客户意见更新{profile['object']}的执行方式，使方案具备持续改进能力。"),
            ("异常处理", f"当出现计划变化、信息不完整或客户侧协调延迟时，及时说明影响范围、调整措施和新的完成时间。"),
            ("交付闭环", f"在每项工作完成后形成提交、确认、归档和复盘动作，保证{profile['evidence']}能够支撑最终验收。"),
        ]
        for title, detail in dimensions:
            paragraphs.append(f"{title}方面，“{subtitle}”将{detail}")

        scenarios = [
            ("常规服务场景", "按照既定计划推进，重点保证服务记录完整、反馈及时、客户确认清楚。"),
            ("重点保障场景", "提高沟通频率和复核强度，对关键事项形成专项说明或阶段报告。"),
            ("异常波动场景", "先控制影响范围，再组织原因分析和恢复验证，最后输出闭环材料。"),
            ("验收准备场景", "按评分项整理材料索引，补齐记录缺口，确保章节内容与附件材料一致。"),
        ]
        for name, detail in scenarios:
            paragraphs.append(f"在{name}下，“{subtitle}”的编写重点有所不同：{detail}")

        return paragraphs

    def long_form_paragraph(self, section_title, subtitle, scoring_text, context, index):
        profile = self.subtitle_profile(section_title, subtitle)
        context_points = self.context_points(context)
        context_point = self.choose_context_point(section_title, subtitle, context_points, index) if context_points else profile["problem"]
        phases = [
            ("准备环节", "完成资料核对、边界确认和计划编排"),
            ("执行环节", "按照服务计划推进具体动作，并同步记录执行过程"),
            ("复核环节", "检查执行结果、材料完整性和客户确认状态"),
            ("反馈环节", "向客户说明完成情况、发现问题和后续建议"),
            ("归档环节", "整理过程证据、报告材料和问题闭环记录"),
        ]
        concerns = [
            ("针对性", f"围绕“{subtitle}”对应的客户场景展开，而不是沿用通用服务说明"),
            ("可行性", f"把{profile['action']}拆成可执行步骤，并明确每一步的触发条件"),
            ("完整性", f"同时覆盖{profile['object']}、过程控制、客户协同和交付证明"),
            ("可验收性", f"将{profile['metric']}写成客户可检查的结果口径"),
            ("连续性", "保证前序分析、服务执行、结果反馈和后续改进之间能够衔接"),
            ("风险控制", f"提前处理{profile['risk']}等可能影响交付质量的因素"),
            ("材料闭环", f"用{profile['evidence']}证明正文描述不是停留在承诺层面"),
        ]
        phase, phase_action = phases[index % len(phases)]
        concern, concern_action = concerns[(index // len(phases)) % len(concerns)]
        return (
            f"从{phase}看，“{subtitle}”需要{phase_action}。"
            f"本段重点体现{concern}：{concern_action}。"
            f"结合“{context_point}”这一项目约束，方案将说明服务方如何组织资源、如何与客户确认、如何输出材料，以及如何把执行结果回写到“{section_title}”的整体交付闭环中。"
        )

    def subtitle_profile(self, section_title, subtitle):
        text = section_title + subtitle
        profiles = [
            (["背景", "整体情况"], ("项目背景与建设现状", "需求理解和关键点分析", "项目理解说明、需求分析记录、关键问题清单", "业务暴露面扩大、现状信息不完整、理解偏差", "需求覆盖度、关键点识别完整度", "客户希望通过本项目补齐 Web 应用防护服务能力并形成统一管理")),
            (["目标"], ("建设目标和服务目标", "目标拆解、指标定义和路径设计", "目标分解表、服务目标矩阵、验收口径说明", "目标空泛、指标不可检查、责任边界不清", "目标可衡量性、目标与评分项对应度", "项目需要把安全效果、服务连续性和管理可视化转化为可检查目标")),
            (["需求"], ("采购需求和评分要求", "需求拆解、优先级排序和响应映射", "需求响应表、评分映射表、章节索引", "需求遗漏、重点偏移、材料无法支撑", "评分项覆盖率、响应完整度", "评分标准要求正文能够体现全面性、针对性和可行性")),
            (["重点", "难点"], ("项目重点难点", "风险识别、难点拆解和应对设计", "重点难点分析表、风险清单、应对措施说明", "策略误配、协同低效、响应不及时、材料缺失", "风险识别完整度、应对措施可执行性", "项目需要提前识别影响交付质量和客户体验的关键因素")),
            (["服务范围", "服务内容"], ("服务范围和服务内容", "服务边界确认、内容编排和交付定义", "服务清单、任务分解表、交付物清单", "范围不清、职责交叉、服务遗漏", "服务覆盖率、交付物完整度", "客户需要清楚看到服务提供方具体做什么、何时做、输出什么")),
            (["服务架构", "组织", "机制"], ("服务组织和协同机制", "角色分工、接口设置和协同流转", "组织架构表、通讯录、职责矩阵、会议纪要", "单点沟通、职责不清、升级不及时", "接口清晰度、协同效率", "服务需要通过稳定组织机制保证跨角色协同")),
            (["流程"], ("服务流程", "流程编排、节点控制和闭环管理", "流程图、节点说明、过程记录", "流程断点、反馈滞后、闭环缺失", "流程完整度、闭环率", "客户关注服务过程是否可跟踪、可复核、可改进")),
            (["进度", "计划", "里程碑"], ("项目进度和里程碑", "计划编制、节点跟踪和偏差纠正", "项目计划表、里程碑清单、周报、延期说明", "计划延期、资源冲突、依赖事项阻塞", "计划完成率、里程碑达成率", "项目需要按阶段推进并保留可检查的进度记录")),
            (["质量"], ("服务质量", "质量目标设定、过程检查和问题改进", "质量检查记录、问题闭环表、客户反馈记录", "质量波动、问题重复、复核不足", "整改闭环率、客户满意度、报告准确率", "客户关注服务质量是否有机制保障，而不只是口头承诺")),
            (["安全", "数据", "文档"], ("信息安全与资料管理", "权限控制、数据保护和文档反馈", "授权记录、资料交接记录、文档版本记录、安全检查记录", "资料泄露、误操作、权限越界、文档失控", "授权合规率、资料归档完整度", "服务过程会接触客户资料和系统信息，必须保证安全可控")),
            (["对接", "沟通", "反馈"], ("服务对接和沟通反馈", "接口确认、沟通组织和成果反馈", "对接通讯录、会议纪要、反馈报告、事项跟踪表", "沟通遗漏、确认不清、反馈不及时", "事项闭环率、反馈及时率", "客户需要在服务过程中持续掌握事项进展和处理结果")),
            (["培训", "课程"], ("培训服务", "课程设计、培训实施和效果评估", "培训计划、课件、签到表、答疑记录、评估表", "培训对象不匹配、内容泛化、效果不可验证", "培训覆盖率、反馈满意度、问题掌握度", "项目需要让客户相关人员理解服务内容和运维协同方式")),
            (["巡检", "维护", "台账"], ("日常维护与巡检", "周期检查、问题发现和台账管理", "巡检记录、问题台账、巡检报告、整改建议", "故障隐患未发现、记录不完整、重复问题未复盘", "巡检覆盖率、问题闭环率", "客户关注服务方能否主动发现风险并持续维护服务状态")),
            (["响应", "故障", "到场", "处置", "闭环"], ("服务响应和故障处置", "事件受理、等级判定、处置恢复和复盘闭环", "工单记录、故障报告、响应记录、复盘报告", "响应超时、影响扩大、原因不清、恢复未验证", "响应及时率、处置完成率、复盘改进率", "服务期间可能出现不同类型故障，需要清晰响应机制")),
        ]
        for keywords, values in profiles:
            if any(keyword in text for keyword in keywords):
                return {
                    "object": values[0],
                    "action": values[1],
                    "evidence": values[2],
                    "risk": values[3],
                    "metric": values[4],
                    "problem": values[5],
                }
        return {
            "object": "本小节服务事项",
            "action": "需求确认、措施执行和结果反馈",
            "evidence": "执行记录、问题清单、阶段报告和确认材料",
            "risk": "需求理解偏差、执行记录缺失和客户反馈滞后",
            "metric": "执行完整度、材料完整度和客户确认情况",
            "problem": "客户需要看到方案内容具备明确对象、执行路径和验收依据",
        }

    def document_text_length(self, document):
        paragraph_chars = sum(len(paragraph.text.strip()) for paragraph in document.paragraphs)
        table_chars = 0
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    table_chars += len(cell.text.strip())
        return paragraph_chars + table_chars

    def section_paragraph_bank(self, section_title, subtitle, scoring_text, context):
        title = section_title + subtitle
        subtitle_focus = self.subtitle_paragraphs(section_title, subtitle)
        scoring_brief = self.scoring_brief(scoring_text)
        common_opening = (
            f"本小节围绕“{subtitle}”展开，重点响应“{section_title}”对应评分要求。"
            f"评分条款重点关注：{scoring_brief}"
        )
        technical_sentence = self.technical_context_sentence(section_title, subtitle, context)

        if "项目整体情况" in title or "关键点分析" in title:
            return [
                common_opening,
                *subtitle_focus,
                technical_sentence,
            ]

        if "整体服务方案" in title:
            return [
                common_opening,
                *subtitle_focus,
                technical_sentence,
            ]

        if "进度" in title or "计划" in title:
            return [
                common_opening,
                *subtitle_focus,
                technical_sentence,
            ]

        if "质量" in title:
            return [
                common_opening,
                *subtitle_focus,
                technical_sentence,
            ]

        if "信息安全" in title or "安全保障" in title:
            return [
                common_opening,
                *subtitle_focus,
                technical_sentence,
            ]

        if "对接" in title:
            return [
                common_opening,
                *subtitle_focus,
                technical_sentence,
            ]

        if "培训" in title:
            return [
                common_opening,
                *subtitle_focus,
                technical_sentence,
            ]

        if "维护" in title or "巡检" in title:
            return [
                common_opening,
                *subtitle_focus,
                technical_sentence,
            ]

        if "响应" in title:
            return [
                common_opening,
                *subtitle_focus,
                technical_sentence,
            ]

        return [
            common_opening,
            *subtitle_focus,
            technical_sentence,
            f"本小节将“{subtitle}”拆解为目标、措施、执行、检查和交付五类内容，确保章节不只停留在原则描述，而是形成可执行工作安排。",
            "具体实施过程中，应先确认需求范围和约束条件，再制定执行计划，并通过过程记录和结果确认保证服务质量。",
            "对应交付物应包括计划、记录、报告、问题清单和确认材料，便于采购方在评审和验收阶段进行核验。",
        ]

    def subtitle_paragraphs(self, section_title, subtitle):
        key = subtitle
        specific = {
            "项目背景理解": [
                "项目背景理解应说明采购方为什么需要补齐 Web 应用安全防护能力，重点体现互联网出口、业务应用暴露面和统一防护管理之间的关系。",
                "方案应从现状约束出发，分析既有网络与未来规划网络对标准化 Web 防护服务的需求，避免把项目理解写成通用产品介绍。",
                "背景分析还需要指出采购服务模式的价值，即通过专业厂商服务提升策略维护、事件响应、巡检分析和持续优化能力。",
            ],
            "建设目标分析": [
                "建设目标应拆解为安全目标、服务目标和管理目标三个层面，分别回应防护能力、服务连续性和过程可视化要求。",
                "安全目标关注攻击拦截、风险发现和策略优化；服务目标关注响应及时、处置闭环和巡检持续；管理目标关注报告完整、材料可查和验收可证。",
                "目标设置需要避免空泛表述，应与后续服务范围、进度计划、质量保障和响应机制建立对应关系。",
            ],
            "关键需求分析": [
                "关键需求分析应把评分表中的高分章节转化为正文编写任务，明确哪些内容必须写深、哪些内容需要用表格或交付物支撑。",
                "从评分条款看，采购方关注的不只是技术能力，还包括服务过程、沟通机制、质量控制、信息安全、培训和维护巡检。",
                "因此本方案需要以评分项为牵引组织章节结构，并在每章中给出可执行措施和可验收材料。",
            ],
            "项目重点分析": [
                "项目重点在于服务方案与采购需求的贴合程度，尤其是 WAF 防护服务如何融入客户现有运维和安全管理流程。",
                "另一个重点是过程材料的完整性，包括计划、记录、报告、台账、培训材料和闭环证明，这些材料直接影响评审和验收。",
                "方案编写应突出持续服务能力，而不是只描述一次性交付动作。",
            ],
            "项目难点分析": [
                "项目难点主要来自需求边界确认、策略调整风险、跨角色沟通、故障场景复杂性和服务过程留痕要求。",
                "对于防护策略相关工作，需要兼顾安全性和业务连续性，避免策略变更造成误拦截或遗漏风险。",
                "对于服务协同相关工作，需要通过例会、工单、升级和确认机制降低沟通成本。",
            ],
            "应对思路": [
                "应对思路应采用“先理解、再规划、后执行、持续复盘”的路径，把项目理解转化为可操作的服务动作。",
                "针对关键服务事项，项目组应建立责任分工、时间节点、质量检查和客户确认机制。",
                "针对可能变化的需求，应保留滚动调整空间，并通过书面记录保证调整过程可追溯。",
            ],
            "服务目标": [
                "服务目标应明确本项目希望达到的运行效果和管理效果，包括防护能力稳定、服务响应及时、巡检维护持续和交付材料完整。",
                "目标描述需要与采购方评分关注点保持一致，既体现技术服务能力，也体现项目管理和客户协同能力。",
                "每项目标都应能够在后续章节中找到对应的执行动作和验收材料。",
            ],
            "服务范围": [
                "服务范围应覆盖项目启动、需求确认、策略梳理、服务实施、日常维护、巡检分析、问题处理、培训支持、阶段汇报和验收配合。",
                "范围边界需要说明哪些事项属于常规服务，哪些事项需要专项确认，避免后续执行时出现职责不清。",
                "对于涉及客户环境的信息、账号、策略和数据，应在范围中明确授权和确认要求。",
            ],
            "服务架构": [
                "服务架构建议采用项目负责人统一协调、技术服务人员负责执行、质量复核人员负责检查、客户接口人负责确认的模式。",
                "该架构能够把服务受理、任务分派、执行反馈和质量复核串联起来，减少单点沟通造成的信息遗漏。",
                "对于重大问题，应在架构中预留升级通道，确保管理层和专家资源能够及时介入。",
            ],
            "服务内容": [
                "服务内容应围绕防护服务运行、策略优化、巡检维护、故障响应、培训支持和报告输出展开。",
                "每类服务内容需要描述触发条件、执行方式、输出材料和客户确认方式。",
                "服务内容不宜只写能力清单，应进一步说明这些能力如何在项目周期内被持续执行。",
            ],
            "服务流程": [
                "服务流程应从任务受理开始，经过需求确认、方案制定、执行实施、结果验证、材料归档和客户反馈等环节。",
                "流程中的每个节点都需要明确输入、输出和责任角色，确保服务过程可以被跟踪。",
                "对于异常事项，应在流程中设置升级、调整和复盘机制。",
            ],
            "服务交付物": [
                "服务交付物应包括启动计划、服务记录、巡检报告、故障处置报告、培训材料、会议纪要、阶段总结和验收支撑材料。",
                "交付物应按照章节和阶段建立清单，便于客户检查服务是否按照约定执行。",
                "对于关键交付物，应保留版本、时间、责任人和确认状态。",
            ],
            "进度总体安排": [
                "进度总体安排应按照项目启动、需求确认、方案细化、服务实施、阶段检查和验收总结分阶段推进。",
                "每个阶段都需要明确开始条件、完成条件、责任人和输出物，避免仅用自然时间描述计划。",
                "总体安排应为后续里程碑、风险控制和异常处理提供统一时间基准。",
            ],
            "进度计划表": [
                "进度计划表应列明阶段、任务、计划时间、责任角色、交付物和依赖条件。",
                "计划表需要支持滚动更新，当客户需求或外部条件变化时，应记录调整原因和影响范围。",
                "对于关键路径任务，应设置提前预警节点，避免临近截止日期才发现延期风险。",
            ],
            "关键里程碑": [
                "关键里程碑包括启动会完成、需求确认完成、服务方案确认、首轮服务执行、阶段巡检反馈、问题整改闭环、培训完成和验收材料提交。",
                "每个里程碑都应对应明确交付物，作为项目进展检查和客户确认的依据。",
                "里程碑管理有助于把长期服务过程拆解为可检查的阶段成果。",
            ],
            "进度保障措施": [
                "进度保障措施包括周计划滚动更新、风险事项提前预警、关键任务专人跟踪、延期事项升级协调和交付物清单化管理。",
                "项目负责人应定期检查计划执行偏差，并针对偏差原因制定纠正措施。",
                "对需要客户配合的事项，应提前提出协调需求，避免等待审批或信息反馈影响整体进度。",
            ],
            "异常情况处理": [
                "异常情况包括需求变更、客户侧审批延迟、重大故障插入、资源冲突和交付物返工等。",
                "出现异常后，应先判断影响范围，再同步调整原因、处理方案和新的完成时间。",
                "所有异常处理过程应形成记录，作为后续复盘和验收说明的依据。",
            ],
            "风险控制": [
                "风险控制应建立风险识别、风险评估、预警触发、处置跟踪和结果复盘机制。",
                "重点风险包括进度延期、服务质量波动、策略变更影响业务、沟通不充分和材料不完整。",
                "项目组应将高风险事项纳入例会跟踪，并在风险解除后更新状态。",
            ],
            "质量保障目标": [
                "质量保障目标应覆盖响应及时性、处置准确性、文档完整性、客户反馈满意度、巡检覆盖率和整改闭环率。",
                "目标设置应具备可检查性，能够通过服务记录、报告和客户确认材料进行验证。",
                "质量目标同时也是后续质量检查和阶段总结的评价依据。",
            ],
            "质量管理机制": [
                "质量管理机制应设置执行自检、项目复核和客户确认三个层级。",
                "执行人员负责服务动作和记录完整，复核人员负责检查内容准确性，客户确认用于校准服务结果。",
                "机制运行过程中应保留检查记录，避免质量管理停留在口头承诺。",
            ],
            "服务要求保障措施": [
                "服务要求保障措施应把评分项中的服务要求拆解成具体动作，并明确执行频次、责任人和输出物。",
                "对于高频服务事项，应采用标准清单减少遗漏；对于复杂事项，应采用专项方案控制质量。",
                "服务保障措施需要与巡检、响应、培训和对接机制形成协同。",
            ],
            "异常处理措施": [
                "异常处理措施应包括异常发现、影响判断、临时处置、原因分析、整改执行和结果验证。",
                "对于影响客户业务或服务质量的异常，应及时升级并同步客户处理进展。",
                "异常处理完成后，应形成闭环记录和改进建议。",
            ],
            "质量检查与改进": [
                "质量检查应定期核对服务记录、巡检结果、问题台账、报告质量和客户反馈。",
                "发现重复问题时，应开展原因复盘，识别流程、配置、沟通或人员安排中的根因。",
                "改进动作应纳入后续服务计划，形成持续优化机制。",
            ],
            "安全保障目标": [
                "安全保障目标是确保服务过程中的客户资料、系统信息、策略配置和故障数据得到受控管理。",
                "目标还包括服务操作可授权、可记录、可追溯，避免未经确认的操作影响业务连续性。",
                "安全目标应贯穿人员管理、数据管理、文档反馈和风险控制全过程。",
            ],
            "团队管理制度": [
                "团队管理制度应明确人员职责、保密要求、权限边界、奖惩机制和工作交接要求。",
                "服务人员应遵守最小知情和最小权限原则，仅在授权范围内接触客户资料和环境信息。",
                "关键服务事项应设置复核要求，降低个人操作失误带来的风险。",
            ],
            "数据安全管理": [
                "数据安全管理应覆盖资料接收、存储、使用、传输、归档和销毁全生命周期。",
                "涉及客户敏感信息的内容应进行脱敏处理，并通过双方认可的渠道进行传递。",
                "数据使用过程应保留记录，确保后续能够追溯资料来源和处理责任。",
            ],
            "文档反馈机制": [
                "文档反馈机制应明确报告、台账、纪要、问题清单等材料的提交频率和确认方式。",
                "重要文档应保留版本号、提交时间、编写人、复核人和客户确认状态。",
                "对于客户反馈意见，应记录处理结果并更新到后续文档中。",
            ],
            "安全风险控制": [
                "安全风险控制应识别资料泄露、误操作、权限滥用、策略误配和沟通不充分等风险。",
                "针对不同风险设置预防措施、监测方式和应急处置路径。",
                "项目结束时应完成资料交接、归档或销毁确认。",
            ],
            "对接组织机制": [
                "对接组织机制应建立双方通讯录和职责表，明确管理、技术、质量和文档等接口人。",
                "项目负责人负责统筹事项流转，技术负责人负责问题研判和处理，质量负责人负责结果复核。",
                "组织机制应支持日常事项处理和重大事项升级两类场景。",
            ],
            "沟通协调方式": [
                "沟通协调方式可采用例会、即时沟通、工单记录、邮件确认和阶段报告相结合的模式。",
                "不同事项应采用不同确认方式，普通事项可快速反馈，关键事项应书面确认。",
                "沟通记录应保留时间、参与人、结论和后续动作。",
            ],
            "服务内容对接": [
                "服务内容对接应逐项确认服务范围、执行频次、交付物格式和客户配合事项。",
                "对于策略调整、巡检维护、培训安排和故障响应等内容，应明确触发条件和确认流程。",
                "对接结果应形成任务清单，作为服务执行依据。",
            ],
            "策略与时间安排": [
                "策略与时间安排应兼顾服务计划和客户业务窗口，避免服务操作影响业务连续性。",
                "涉及变更或重要服务动作时，应提前沟通时间、影响范围和回退措施。",
                "时间安排发生变化时，应同步调整原因和新的计划。",
            ],
            "成果反馈机制": [
                "成果反馈应包含完成情况、发现问题、处理建议、后续计划和需客户协调事项。",
                "反馈形式可包括服务报告、会议纪要、问题台账和专项说明。",
                "对于客户反馈意见，应明确处理责任和完成时间。",
            ],
            "培训目标": [
                "培训目标是帮助客户相关人员理解服务内容、掌握常见操作、熟悉问题反馈流程并提升安全意识。",
                "培训不只面向系统使用，也应覆盖服务协同、报告解读和日常维护注意事项。",
                "培训目标应与后续培训课程、频次和效果评估保持对应。",
            ],
            "培训对象": [
                "培训对象可分为管理人员、技术运维人员和一线使用人员。",
                "管理人员重点了解服务范围、管理机制和报告结论；运维人员重点掌握巡检、响应和故障报送；使用人员重点了解常见问题和安全注意事项。",
                "不同对象应采用不同深度和案例。",
            ],
            "培训课程安排": [
                "培训课程安排应结合项目启动、运行和验收阶段进行。",
                "启动阶段开展基础培训，运行阶段开展专题培训，验收前开展总结培训和材料确认。",
                "课程安排应提前通知客户，并保留签到和培训材料。",
            ],
            "培训内容设置": [
                "培训内容应包括项目服务范围、日常巡检关注点、故障报送方式、响应流程、常见问题处置、报告解读和安全注意事项。",
                "内容设置应结合客户实际使用场景，避免只讲产品功能。",
                "对于常见故障和高频问题，应配套案例讲解和答疑。",
            ],
            "培训频次": [
                "培训频次可采用项目阶段培训、月度专题培训和按需补充培训相结合的方式。",
                "当服务内容、策略或客户人员发生变化时，应及时补充培训。",
                "培训频次应兼顾覆盖率和客户工作安排。",
            ],
            "培训效果评估": [
                "培训效果评估应通过签到记录、问题答疑、现场反馈、问卷结果和后续使用情况进行验证。",
                "对于培训中发现的理解偏差，应补充说明材料或安排专项答疑。",
                "评估结果应纳入阶段服务总结。",
            ],
            "维护目标": [
                "维护目标是保障服务状态稳定、问题及时发现、风险提前预警和服务材料持续完整。",
                "维护工作应从被动处理转向主动发现，通过巡检和趋势分析降低故障发生概率。",
                "维护目标需要与响应机制和质量保障机制形成联动。",
            ],
            "巡检频次": [
                "巡检频次可根据项目要求设置为日常检查、周度汇总、月度分析和专项检查。",
                "重大活动、策略调整或高风险阶段可临时提高巡检频率。",
                "巡检频次调整应说明原因并同步客户确认。",
            ],
            "巡检内容": [
                "巡检内容应覆盖服务运行状态、策略配置、告警事件、关键指标、问题工单、历史故障和客户反馈事项。",
                "每项巡检内容应明确检查方法、结果判定和异常处理方式。",
                "巡检内容应形成标准清单，保证不同人员执行时结果一致。",
            ],
            "巡检台账": [
                "巡检台账应记录巡检时间、检查项、检查结果、异常说明、处理建议、责任人和闭环状态。",
                "台账既是过程管理材料，也是后续趋势分析和验收支撑材料。",
                "对于重复出现的问题，应在台账中标记并纳入专项分析。",
            ],
            "异常反馈机制": [
                "异常反馈机制应规定异常发现后的通知方式、反馈时限、影响说明和处理路径。",
                "重大异常应及时升级，普通异常应纳入问题台账持续跟踪。",
                "反馈内容应避免只写现象，还应说明建议措施和后续计划。",
            ],
            "故障预防": [
                "故障预防应基于巡检结果、历史问题和客户反馈识别易发故障点。",
                "对高风险事项应提前提出优化建议或专项检查计划。",
                "预防措施应定期复盘效果，持续降低重复问题发生概率。",
            ],
            "响应目标": [
                "响应目标是确保服务事件能够被快速受理、准确分级、及时处置和完整闭环。",
                "目标应覆盖响应时间、到场机制、处置效率、客户反馈和报告归档。",
                "响应目标需要与客户服务要求保持一致。",
            ],
            "故障场景分类": [
                "故障场景可根据影响范围、紧急程度和业务影响分为重大、紧急、一般和咨询类事项。",
                "不同级别应对应不同的响应动作、人员投入和升级要求。",
                "场景分类有助于避免所有问题都按同一方式处理。",
            ],
            "响应时间": [
                "响应时间应在接到通知后及时确认事项、记录现象、判断影响范围并启动处理流程。",
                "对于重大和紧急事项，应优先完成影响控制和客户同步。",
                "响应时间记录应进入服务台账，作为服务质量评价依据。",
            ],
            "人员到场机制": [
                "人员到场机制应明确远程支持和现场支持的触发条件。",
                "当问题影响范围较大、远程无法定位或客户要求现场协同时，应组织具备相应能力的人员到场。",
                "到场人员应携带必要资料和处置计划，并在完成后提交处理记录。",
            ],
            "故障处置流程": [
                "故障处置流程应包括受理登记、初步研判、临时止损、原因定位、修复处理、结果验证、客户确认和报告归档。",
                "流程执行中应持续同步处置进展，避免客户无法掌握处理状态。",
                "处置完成后，应确认影响是否消除，并记录根因和改进建议。",
            ],
            "升级与闭环": [
                "升级与闭环机制适用于超时、影响扩大、原因不明或需要客户协调的问题。",
                "升级后应明确新的负责人、处理计划和反馈频率。",
                "闭环时应形成处置报告、客户确认和复盘改进措施。",
            ],
        }
        if key in specific:
            return specific[key]
        return [
            f"“{subtitle}”应围绕“{section_title}”的章节目标展开，明确本小节的服务对象、执行动作和输出材料。",
            "编写时应优先回应评分标准中的关键词，并结合技术规范书补充项目约束和客户场景。",
            "本小节内容应形成可执行、可检查、可归档的工作说明。",
        ]

    def add_section_table(self, document, section_title):
        presets = [
            ("进度", ["阶段", "工作内容", "输出物", "保障措施"], [
                ["启动准备", "召开启动会，确认项目范围、接口人和关键节点", "启动纪要、项目计划", "项目负责人统一跟踪"],
                ["方案细化", "结合评分标准和技术规范书细化执行方案", "细化方案、任务清单", "关键内容双人复核"],
                ["服务实施", "按计划开展服务、巡检、培训和响应工作", "服务记录、巡检报告", "周计划滚动更新"],
                ["验收总结", "汇总交付物并配合客户完成确认", "验收材料、总结报告", "问题闭环后提交"],
            ]),
            ("质量", ["质量环节", "控制动作", "输出记录", "责任角色"], [
                ["执行前", "确认需求、范围、计划和交付标准", "需求确认表", "项目负责人"],
                ["执行中", "检查服务动作、文档记录和问题处理状态", "过程检查记录", "质量复核人员"],
                ["执行后", "复核结果并收集客户反馈", "服务报告、反馈记录", "服务负责人"],
                ["改进", "对重复问题进行复盘并优化流程", "改进清单", "项目组"],
            ]),
            ("信息安全", ["控制域", "保障措施", "过程记录", "责任角色"], [
                ["人员管理", "落实保密要求和权限边界", "人员授权记录", "项目负责人"],
                ["数据保护", "敏感信息分级、脱敏和受控传输", "资料交接记录", "服务人员"],
                ["操作安全", "执行前确认授权，执行后记录结果", "操作记录", "技术负责人"],
                ["文档归档", "保留版本、时间和确认状态", "归档清单", "文档负责人"],
            ]),
            ("对接", ["对接事项", "沟通方式", "反馈内容", "升级条件"], [
                ["日常事项", "例会、工单、即时沟通", "完成情况、待协调事项", "连续延期或需跨部门协调"],
                ["服务变更", "书面确认、会议评审", "变更范围、影响分析", "影响进度或服务质量"],
                ["重大问题", "电话通知、专项会议", "影响范围、处置进展", "超时未恢复或影响扩大"],
            ]),
            ("培训", ["培训对象", "课程内容", "组织方式", "效果验证"], [
                ["管理人员", "服务范围、管理机制、报告解读", "集中培训", "反馈记录"],
                ["运维人员", "巡检内容、故障报送、响应流程", "专题培训", "答疑记录"],
                ["使用人员", "常见问题、安全注意事项", "线上或现场培训", "签到与问卷"],
            ]),
            ("巡检", ["巡检对象", "巡检内容", "频次", "输出材料"], [
                ["服务状态", "运行状态、关键指标、异常告警", "日常/周度", "巡检记录"],
                ["问题工单", "处理进展、闭环状态、重复问题", "周度", "问题台账"],
                ["趋势分析", "风险趋势、优化建议、阶段总结", "月度", "巡检报告"],
            ]),
            ("响应", ["事件级别", "适用场景", "响应动作", "闭环要求"], [
                ["重大事件", "影响核心业务或范围较大", "立即确认并启动升级处置", "形成专项报告和复盘"],
                ["紧急事件", "影响部分业务或存在扩大风险", "快速研判并组织处理", "记录原因和修复结果"],
                ["一般事件", "常规问题、咨询或优化需求", "按流程受理并反馈", "工单闭环和材料归档"],
            ]),
        ]

        for keyword, headers, rows in presets:
            if keyword in section_title:
                table = document.add_table(rows=1, cols=len(headers))
                table.style = "Table Grid"
                for index, header in enumerate(headers):
                    table.rows[0].cells[index].text = header
                for row in rows:
                    cells = table.add_row().cells
                    for index, value in enumerate(row):
                        cells[index].text = value
                document.add_paragraph("")
                return

    def has_section_table(self, section_title):
        return any(keyword in section_title for keyword in ["进度", "质量", "信息安全", "对接", "培训", "巡检", "响应"])

    def technical_context_sentence(self, section_title, subtitle, context):
        points = self.context_points(context)
        if not points:
            return f"由于技术规范书未提取到可直接引用的细化条款，“{subtitle}”将以评分标准为主线展开，同时在正文中预留与技术规范书核对的位置，便于后续补充项目专属参数、服务边界和验收口径。"
        selected = self.choose_context_point(section_title, subtitle, points, 0)
        return f"结合技术规范书中与“{subtitle}”最相关的要求，本小节重点吸收“{selected}”这一信息，并将其转化为服务动作、过程记录和交付材料。"

    def context_points(self, context):
        if not context or "结合技术规范书要求进行内容补充" in context:
            return []
        parts = re.split(r"[；;。\\n]+", context)
        return [self.clean_text(part) for part in parts if len(self.clean_text(part)) >= 12][:30]

    def context_summary(self, technical_signals):
        cleaned = [self.clean_text(signal) for signal in technical_signals if self.clean_text(signal)]
        if not cleaned:
            return "结合技术规范书要求进行内容补充，并以评分标准作为章节展开依据。"
        return "；".join(cleaned[:12])

    def choose_context_point(self, section_title, subtitle, points, index=0):
        keywords = self.context_keywords(section_title + subtitle)
        scored = []
        for point_index, point in enumerate(points):
            score = sum(1 for keyword in keywords if keyword and keyword in point)
            scored.append((score, -point_index, point))
        scored.sort(reverse=True)
        matched = [point for score, _, point in scored if score > 0]
        if matched:
            return matched[index % len(matched)]
        return points[index % len(points)]

    def context_keywords(self, text):
        mapping = [
            (["背景", "整体情况", "关键"], ["Web应用防护", "统一", "互联网出口", "建设", "防护能力"]),
            (["部署", "实施", "进度", "计划"], ["部署", "安装", "调测", "观察模式", "阻断模式", "验证"]),
            (["服务", "整体", "范围", "内容"], ["服务", "Web应用防护", "漏洞应急", "规则", "版本升级"]),
            (["质量", "保障"], ["误报率", "验证", "报告", "灰度", "规则", "调优"]),
            (["安全", "数据", "信息"], ["HTTPS", "证书", "解密", "安全", "权限", "日志"]),
            (["培训"], ["培训", "管理员", "安全分析", "运维人员", "实操", "演练"]),
            (["巡检", "维护", "台账"], ["巡检", "维护", "运维", "告警", "报表", "规则库", "版本升级"]),
            (["响应", "故障", "到场", "处置"], ["漏洞应急", "12小时", "24小时", "故障", "响应", "处置", "告警"]),
            (["对接", "沟通", "反馈"], ["采购人", "提供", "报告", "反馈", "服务"]),
        ]
        keywords = []
        for triggers, values in mapping:
            if any(trigger in text for trigger in triggers):
                keywords.extend(values)
        return keywords or ["服务", "要求", "支持"]

    def clean_text(self, text):
        text = re.sub(r"\s+", " ", str(text or "")).strip()
        text = text.replace("|", "，")
        return text[:320] + ("..." if len(text) > 320 else "")

    def scoring_brief(self, text):
        text = self.clean_text(text)
        text = re.sub(r"注：.*$", "", text).strip()
        text = re.sub(r"根据.*?进行评审[，,]?", "", text).strip()
        parts = re.split(r"[。；;]", text)
        brief = parts[0] if parts else text
        brief = brief.replace("响应供应商", "供应商")
        brief = re.sub(r"（?\\d+分）?", "", brief)
        brief = re.sub(r"得[^。；;]{0,30}分", "", brief)
        return brief[:110] + ("..." if len(brief) > 110 else "")

    def parse_multipart_form(self):
        content_type = self.headers.get("Content-Type", "")
        if not content_type.startswith("multipart/form-data"):
            raise ValueError("请求必须使用 multipart/form-data")

        content_length = int(self.headers.get("Content-Length", "0") or 0)
        if content_length <= 0:
            raise ValueError("上传请求缺少有效的 Content-Length")
        if content_length > MAX_MULTIPART_BYTES:
            raise ValueError("上传文件过大，请压缩后重试")

        return cgi.FieldStorage(
            fp=self.rfile,
            headers=self.headers,
            environ={
                "REQUEST_METHOD": "POST",
                "CONTENT_TYPE": content_type,
            },
        )

    def parse_config(self, form):
        raw_config = form.getfirst("config")
        if not raw_config:
            raise ValueError("缺少生成配置 config")

        try:
            return json.loads(raw_config)
        except json.JSONDecodeError as exc:
            raise ValueError(f"config 不是合法 JSON：{exc}") from exc

    def validate_required_files(self, form):
        required = {
            "technicalSpecification": "技术规范书",
            "scoringCriteria": "评分标准",
        }

        for field, label in required.items():
            if field not in form or not getattr(form[field], "filename", ""):
                raise ValueError(f"缺少必传文件：{label}")

    def validate_upload_name(self, field_name, filename):
        suffix = Path(filename or "").suffix.lower()
        allowed = ALLOWED_UPLOAD_SUFFIXES.get(field_name, set())
        if suffix not in allowed:
            raise ValueError(f"{field_name} 文件类型不受支持")

    def save_upload(self, field_name, item, upload_dir):
        self.validate_upload_name(field_name, item.filename)
        filename = safe_filename(item.filename)
        target = upload_dir / filename
        written = 0
        try:
            with target.open("wb") as output:
                while True:
                    chunk = item.file.read(64 * 1024)
                    if not chunk:
                        break
                    written += len(chunk)
                    if written > MAX_UPLOAD_BYTES:
                        raise ValueError("单个上传文件超过大小限制")
                    output.write(chunk)
        except Exception:
            target.unlink(missing_ok=True)
            raise

        return {
            "filename": filename,
            "path": str(target.relative_to(BASE_DIR)),
            "size": written,
        }

    def copy_default_template(self, job_dir):
        target = job_dir / DEFAULT_TEMPLATE.name
        shutil.copy2(DEFAULT_TEMPLATE, target)

        style_target = job_dir / DEFAULT_STYLE.name
        shutil.copy2(DEFAULT_STYLE, style_target)

        return {
            "filename": DEFAULT_TEMPLATE.name,
            "path": str(target.relative_to(BASE_DIR)),
            "styleProfile": str(style_target.relative_to(BASE_DIR)),
            "size": target.stat().st_size,
        }

    def serve_output_file(self):
        relative = unquote(self.path.removeprefix("/output/"))
        relative_path = Path(relative)
        if any(part in DISALLOWED_FILE_SEGMENTS for part in relative_path.parts):
            self.send_error(HTTPStatus.NOT_FOUND, "Output file not found")
            return
        target = (OUTPUT_DIR / relative).resolve()
        if not str(target).startswith(str(OUTPUT_DIR.resolve())) or not target.is_file():
            self.send_error(HTTPStatus.NOT_FOUND, "Output file not found")
            return

        job_id = relative_path.parts[0] if relative_path.parts else ""
        job_dir = (OUTPUT_DIR / safe_job_id(job_id)).resolve()
        self.authorize_job_access(job_dir)

        self.path = "/output/" + relative
        self.directory = str(BASE_DIR)
        super().do_GET()

    def job_belongs_to_session(self, job_dir, session_id):
        try:
            owner_session_id = self.owner_session_id_for_job(job_dir)
        except (OSError, ValueError, json.JSONDecodeError):
            return False
        return owner_session_id == session_id

    def authorize_job_access(self, job_dir, manifest=None):
        session_id = self.require_session()
        owner_session_id = self.owner_session_id_for_job(job_dir, manifest=manifest)
        if owner_session_id != session_id:
            raise PermissionError("当前会话无权访问该任务")

    def owner_session_id_for_job(self, job_dir, manifest=None):
        access_path = job_dir / ACCESS_CONTROL_FILE
        if access_path.exists():
            data = json.loads(access_path.read_text(encoding="utf-8"))
            owner_session_id = data.get("ownerSessionId", "")
            if owner_session_id:
                return owner_session_id
        if manifest is None:
            manifest_path = job_dir / "manifest.json"
            if not manifest_path.exists():
                raise ValueError("任务清单不存在")
            manifest = json.loads(manifest_path.read_text(encoding="utf-8"))
        return manifest.get("ownerSessionId", "")

    def send_json(self, payload, status):
        body = json.dumps(payload, ensure_ascii=False, indent=2).encode("utf-8")
        self.send_response(status)
        self.send_header("Content-Type", "application/json; charset=utf-8")
        self.send_header("Content-Length", str(len(body)))
        self.end_headers()
        self.wfile.write(body)


def safe_filename(filename):
    keep = []
    for char in Path(filename).name:
        if char.isalnum() or char in {".", "-", "_", " ", "（", "）", "(", ")"}:
            keep.append(char)
    cleaned = "".join(keep).strip()
    return cleaned or f"upload-{uuid.uuid4().hex}"


def safe_job_id(job_id):
    keep = []
    for char in job_id:
        if char.isalnum() or char in {"-", "_"}:
            keep.append(char)
    cleaned = "".join(keep).strip()
    if not cleaned:
        raise ValueError("任务编号无效")
    return cleaned


def main():
    logging.basicConfig(level=logging.INFO, format="%(asctime)s %(levelname)s %(message)s")
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    server = ThreadingHTTPServer(("127.0.0.1", 8011), ProposalHandler)
    print("Bid Proposal Agent running at http://127.0.0.1:8011")
    server.serve_forever()


if __name__ == "__main__":
    main()
